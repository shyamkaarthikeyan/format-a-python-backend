#!/usr/bin/env python3
"""
Debug script to isolate the table creation issue.
This will create a minimal document with just a table to see if it appears in Word.
"""

import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_simple_table_test():
    """Create a simple document with just a table to test visibility."""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Simple Table Test")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some text before table
    doc.add_paragraph("This text should appear before the table.")
    
    # Add table caption
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run("TABLE 1: TEST TABLE CAPTION")
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(9)
    caption_run.bold = True
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a simple 3x3 table
    print("Creating simple table...", file=sys.stderr)
    table = doc.add_table(rows=3, cols=3)
    print(f"Table created with {len(table.rows)} rows and {len(table.columns)} columns", file=sys.stderr)
    
    # Set table style
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add very strong borders to make table visible
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Set cell borders to make table visible
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            
            # Add thick black borders
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")  # Thick border (1.5pt)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")  # Black border
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
            
            # Add content to cell
            if row_idx == 0:
                # Header row
                cell.text = f"Header {col_idx + 1}"
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
            else:
                # Data row
                cell.text = f"Data {row_idx},{col_idx + 1}"
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
    
    print("Table content added", file=sys.stderr)
    
    # Add some text after table
    doc.add_paragraph("This text should appear after the table.")
    
    return doc

def main():
    """Create and save the simple table test."""
    print("=== SIMPLE TABLE TEST ===")
    
    try:
        # Create the document
        doc = create_simple_table_test()
        
        # Save to file
        output_file = "debug_table_issue_output.docx"
        doc.save(output_file)
        
        print(f"✅ Simple table test document created: {output_file}")
        print("\n=== EXPECTED IN DOCUMENT ===")
        print("1. Title: 'Simple Table Test'")
        print("2. Text: 'This text should appear before the table.'")
        print("3. TABLE 1: TEST TABLE CAPTION")
        print("4. A 3x3 table with thick black borders containing:")
        print("   - Header row: Header 1, Header 2, Header 3")
        print("   - Data row 1: Data 1,1, Data 1,2, Data 1,3")
        print("   - Data row 2: Data 2,1, Data 2,2, Data 2,3")
        print("5. Text: 'This text should appear after the table.'")
        
        print("\n=== CRITICAL TEST ===")
        print("🔍 Open the document in Microsoft Word and check:")
        print("   ❓ Does the table appear after the caption?")
        print("   ❓ Are the table borders visible?")
        print("   ❓ Is the table content (headers and data) visible?")
        print("   ❓ Is the table properly positioned between the text blocks?")
        
        print("\n=== IF TABLE DOESN'T APPEAR ===")
        print("This indicates a fundamental issue with table creation in python-docx")
        print("Possible causes:")
        print("- Document structure corruption")
        print("- Table positioning issues")
        print("- OpenXML formatting conflicts")
        print("- Word compatibility issues")
        
    except Exception as e:
        print(f"❌ Error creating simple table test: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0

if __name__ == "__main__":
    sys.exit(main())