"""
LaTeX to Word Equation Converter
Converts LaTeX equations to Word OMML (Office Math Markup Language) format
"""

from latex2mathml.converter import convert as latex_to_mathml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


def mathml_to_omml(mathml_str):
    """
    Convert MathML to OMML (Office Math Markup Language) for Word
    This is a simplified converter for basic equations
    """
    # Remove XML declaration and namespace if present
    mathml_str = re.sub(r'<\?xml[^>]*\?>', '', mathml_str)
    mathml_str = re.sub(r'xmlns="[^"]*"', '', mathml_str)
    
    # Basic MathML to OMML mapping
    # This is a simplified version - full conversion would require XSLT
    omml_str = mathml_str
    
    # Replace MathML tags with OMML equivalents
    replacements = {
        '<math>': '<m:oMathPara><m:oMath>',
        '</math>': '</m:oMath></m:oMathPara>',
        '<mrow>': '<m:r>',
        '</mrow>': '</m:r>',
        '<mi>': '<m:r><m:t>',
        '</mi>': '</m:t></m:r>',
        '<mn>': '<m:r><m:t>',
        '</mn>': '</m:t></m:r>',
        '<mo>': '<m:r><m:t>',
        '</mo>': '</m:t></m:r>',
        '<mfrac>': '<m:f><m:fPr></m:fPr>',
        '</mfrac>': '</m:f>',
        '<msqrt>': '<m:rad><m:radPr></m:radPr>',
        '</msqrt>': '</m:rad>',
        '<msup>': '<m:sSup><m:sSupPr></m:sSupPr>',
        '</msup>': '</m:sSup>',
        '<msub>': '<m:sSub><m:sSubPr></m:sSubPr>',
        '</msub>': '</m:sSub>',
    }
    
    for old, new in replacements.items():
        omml_str = omml_str.replace(old, new)
    
    return omml_str


def insert_latex_equation(paragraph, latex_code, equation_number=None):
    """
    Insert a LaTeX equation into a Word paragraph
    
    Args:
        paragraph: python-docx paragraph object
        latex_code: LaTeX equation string
        equation_number: Optional equation number for display
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Convert LaTeX to MathML
        mathml = latex_to_mathml(latex_code)
        
        # For now, we'll insert the equation as formatted text
        # Full OMML conversion requires more complex XSLT transformation
        
        # Center the equation
        paragraph.alignment = 1  # Center alignment
        
        # Add the equation as a run with special formatting
        run = paragraph.add_run()
        
        # Try to render as close to equation as possible
        # This is a fallback - ideally we'd use OMML
        formatted_eq = format_latex_for_display(latex_code)
        run.text = formatted_eq
        run.font.italic = True
        run.font.size = Pt(11)
        
        # Add equation number if provided
        if equation_number is not None:
            number_run = paragraph.add_run(f"  ({equation_number})")
            number_run.font.size = Pt(10)
        
        return True
        
    except Exception as e:
        print(f"Error converting LaTeX equation: {e}", file=sys.stderr)
        # Fallback: insert as plain text
        paragraph.add_run(latex_code)
        if equation_number is not None:
            paragraph.add_run(f"  ({equation_number})")
        return False


def format_latex_for_display(latex_code):
    """
    Format LaTeX code for better display in Word
    Converts common LaTeX commands to Unicode equivalents
    """
    # Common LaTeX to Unicode mappings
    replacements = {
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\delta': 'δ',
        r'\epsilon': 'ε',
        r'\theta': 'θ',
        r'\lambda': 'λ',
        r'\mu': 'μ',
        r'\pi': 'π',
        r'\sigma': 'σ',
        r'\phi': 'φ',
        r'\omega': 'ω',
        r'\Delta': 'Δ',
        r'\Sigma': 'Σ',
        r'\Omega': 'Ω',
        r'\infty': '∞',
        r'\pm': '±',
        r'\times': '×',
        r'\div': '÷',
        r'\leq': '≤',
        r'\geq': '≥',
        r'\neq': '≠',
        r'\approx': '≈',
        r'\equiv': '≡',
        r'\sum': '∑',
        r'\int': '∫',
        r'\partial': '∂',
        r'\nabla': '∇',
        r'\sqrt': '√',
    }
    
    result = latex_code
    for latex, unicode_char in replacements.items():
        result = result.replace(latex, unicode_char)
    
    # Handle fractions: \frac{a}{b} -> a/b
    result = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', result)
    
    # Handle superscripts: x^{2} -> x²
    superscripts = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', 
                   '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
                   'n': 'ⁿ', 'i': 'ⁱ'}
    result = re.sub(r'\^{([0-9ni])}', lambda m: superscripts.get(m.group(1), '^' + m.group(1)), result)
    result = re.sub(r'\^([0-9ni])', lambda m: superscripts.get(m.group(1), '^' + m.group(1)), result)
    
    # Handle subscripts: x_{i} -> xᵢ
    subscripts = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
                 '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
                 'i': 'ᵢ', 'j': 'ⱼ', 'n': 'ₙ'}
    result = re.sub(r'_{([0-9ijn])}', lambda m: subscripts.get(m.group(1), '_' + m.group(1)), result)
    result = re.sub(r'_([0-9ijn])', lambda m: subscripts.get(m.group(1), '_' + m.group(1)), result)
    
    # Remove remaining braces
    result = result.replace('{', '').replace('}', '')
    
    # Remove backslashes from remaining commands
    result = re.sub(r'\\([a-zA-Z]+)', r'\1', result)
    
    return result


# Import required for font size
from docx.shared import Pt
import sys
