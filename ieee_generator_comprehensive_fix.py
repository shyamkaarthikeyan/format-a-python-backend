"""
IEEE Document Generator - Comprehensive Fix for PDF Justification, Image Positioning, and Table Names
"""

import json
import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re
from html.parser import HTMLParser
import unicodedata
import base64
import tempfile
import subprocess

def sanitize_text(text):
    """Sanitize text to remove invalid Unicode characters and surrogates."""
    if not text:
        return ""
    
    # Convert to string if not already
    text = str(text)
    
    # Remove surrogate characters and other problematic Unicode
    text = text.encode('utf-8', 'ignore').decode('utf-8')
    
    # Normalize Unicode characters
    text = unicodedata.normalize('NFKD', text)
    
    # Remove any remaining control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    return text

# IEEE EXACT LATEX PDF FORMATTING - LOW-LEVEL OPENXML SPECIFICATIONS
IEEE_CONFIG = {
    'font_name': 'Times New Roman',
    # Font sizes (exact IEEE LaTeX specifications)
    'font_size_title': Pt(24),        # Title: 24pt bold centered
    'font_size_author_name': Pt(10),  # Author names: 10pt bold
    'font_size_author_affil': Pt(10), # Author affiliations: 10pt italic
    'font_size_author_email': Pt(9),  # Author emails: 9pt
    'font_size_body': Pt(10),         # Body text: 10pt
    'font_size_abstract': Pt(9),      # Abstract/Keywords: 9pt bold
    'font_size_caption': Pt(9),       # Captions: 9pt italic
    'font_size_reference': Pt(9),     # References: 9pt
    
    # Page margins (exact IEEE LaTeX: 0.75" all sides = 1080 twips)
    'margin_twips': 1080,
    
    # Two-column layout (exact IEEE LaTeX specifications)
    'column_count': 2,
    'column_width_twips': 4770,       # 3.3125" per column
    'column_gap_twips': 360,          # 0.25" gap between columns
    'column_indent': Pt(0),           # No indent for column text
    
    # Line spacing (exact IEEE LaTeX: 12pt = 240 twips)
    'line_spacing_twips': 240,
    'line_spacing': Pt(12),  # For backward compatibility
    
    # Paragraph spacing (exact IEEE LaTeX specifications)
    'spacing_title_after': 240,       # 12pt after title
    'spacing_abstract_after': 120,    # 6pt after abstract
    'spacing_keywords_after': 240,    # 12pt after keywords
    'spacing_section_before': 240,    # 12pt before section headings
    'spacing_section_after': 0,       # 0pt after section headings
    
    # Figure specifications
    'figure_max_width_twips': 4770,   # Max 3.3125" width (column width)
    'figure_spacing': 120,             # 6pt before/after figures
    'figure_sizes': {
        'Very Small': Inches(1.5),     # 1.5" width
        'Small': Inches(2.0),          # 2.0" width  
        'Medium': Inches(2.5),         # 2.5" width
        'Large': Inches(3.3125)        # Full column width
    },
    'max_figure_height': Inches(4.0),  # Max figure height
    
    # Reference specifications
    'reference_hanging_indent': 360,  # 0.25" hanging indent
}

def apply_perfect_justification(para):
    """Apply perfect text justification for both DOCX and PDF output."""
    # Set paragraph alignment to justified
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Apply OpenXML justification settings
    pPr = para._element.get_or_add_pPr()
    
    # Clear existing alignment
    for elem in pPr.xpath('./w:jc'):
        pPr.remove(elem)
    
    # Set full justification
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')  # Full justification
    pPr.append(jc)
    
    # Add text distribution for better spacing
    textAlignment = OxmlElement('w:textAlignment')
    textAlignment.set(qn('w:val'), 'distribute')
    pPr.append(textAlignment)
    
    # Enable automatic hyphenation for better line breaks
    suppressAutoHyphens = OxmlElement('w:suppressAutoHyphens')
    suppressAutoHyphens.set(qn('w:val'), '0')  # Enable hyphenation
    pPr.append(suppressAutoHyphens)

def add_image_with_proper_positioning(doc, image_data, width, caption=None, figure_num=None):
    """Add image with proper positioning to prevent text overlap in DOCX."""
    try:
        # Decode base64 image data
        if ',' in image_data:
            image_data = image_data.split(',')[1]
        
        image_bytes = base64.b64decode(image_data)
        image_stream = BytesIO(image_bytes)
        
        # Create a new paragraph for the image with proper spacing
        image_para = doc.add_paragraph()
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing before image
        image_para.paragraph_format.space_before = Pt(12)
        image_para.paragraph_format.space_after = Pt(6)
        
        # Add the image
        run = image_para.add_run()
        picture = run.add_picture(image_stream, width=width)
        
        # Scale if too tall
        if picture.height > IEEE_CONFIG['max_figure_height']:
            scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
            run.clear()
            image_stream.seek(0)
            run.add_picture(image_stream, width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
        
        # Add caption if provided
        if caption and figure_num:
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(f"FIG. {figure_num}: {sanitize_text(caption).upper()}")
            caption_run.font.name = 'Times New Roman'
            caption_run.font.size = Pt(9)
            caption_run.bold = True
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.paragraph_format.space_before = Pt(3)
            caption_para.paragraph_format.space_after = Pt(12)
        
        # Add spacing after image block
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(6)
        
        return True
        
    except Exception as e:
        print(f"Error adding image: {e}", file=sys.stderr)
        return False

def get_table_caption(table_data):
    """Get table caption, avoiding duplication between tableName and caption."""
    # Priority: caption > tableName > default
    caption = table_data.get('caption', '').strip()
    table_name = table_data.get('tableName', '').strip()
    
    # If both exist and are different, use caption
    if caption and table_name:
        # Check if caption already contains table name or vice versa
        if table_name.lower() in caption.lower():
            return caption
        elif caption.lower() in table_name.lower():
            return table_name
        else:
            # They're different, prefer caption
            return caption
    
    # Use whichever exists
    return caption or table_name or 'Data Table'

def add_ieee_table_fixed(doc, table_data, section_idx, table_count):
    """Add table with fixed caption handling to prevent duplication."""
    try:
        table_type = table_data.get('tableType', table_data.get('type', 'interactive'))
        
        # Get the table caption (avoiding duplication)
        table_caption = get_table_caption(table_data)
        
        # Add table caption BEFORE the table
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(f"TABLE {section_idx}.{table_count}: {sanitize_text(table_caption).upper()}")
        caption_run.font.name = 'Times New Roman'
        caption_run.font.size = Pt(9)
        caption_run.bold = True
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.paragraph_format.space_before = Pt(12)
        caption_para.paragraph_format.space_after = Pt(6)
        
        if table_type == 'interactive':
            # Handle interactive tables
            headers = table_data.get('headers', [])
            rows_data = table_data.get('tableData', [])
            
            if not headers or not rows_data:
                print(f"Warning: Interactive table missing headers or data", file=sys.stderr)
                return
            
            # Create table
            num_cols = len(headers)
            num_rows = len(rows_data) + 1  # +1 for header row
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set column widths
            col_width = 4770 // num_cols
            for col in table.columns:
                col.width = col_width
            
            # Header row
            header_row = table.rows[0]
            for col_idx, header in enumerate(headers):
                cell = header_row.cells[col_idx]
                cell.text = sanitize_text(str(header))
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                        run.bold = True
            
            # Data rows
            for row_idx, row_data in enumerate(rows_data):
                table_row = table.rows[row_idx + 1]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table_row.cells[col_idx]
                        cell.text = sanitize_text(str(cell_data))
                        
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                                run.bold = False
        
        elif table_type == 'image':
            # Handle image tables with proper positioning
            if table_data.get('data'):
                size = table_data.get('size', 'medium')
                size_mapping = {
                    'small': Inches(2.0),
                    'medium': Inches(3.0),
                    'large': Inches(3.3125)
                }
                width = size_mapping.get(size, Inches(3.0))
                
                # Use the improved image positioning function
                add_image_with_proper_positioning(doc, table_data['data'], width)
        
        # Add spacing after table
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(6)
        spacing_para.paragraph_format.space_after = Pt(12)
        
    except Exception as e:
        print(f"Error adding table: {e}", file=sys.stderr)

def generate_pdf_with_perfect_justification(html_content, output_path):
    """Generate PDF with perfect text justification using WeasyPrint."""
    
    # Enhanced CSS for perfect PDF justification
    enhanced_css = """
    <style>
    @page {
        size: A4;
        margin: 0.75in;
    }
    
    body {
        font-family: 'Times New Roman', serif;
        font-size: 10pt;
        line-height: 1.2;
        text-align: justify !important;
        text-justify: inter-word !important;
        hyphens: auto !important;
        -webkit-hyphens: auto !important;
        -moz-hyphens: auto !important;
        -ms-hyphens: auto !important;
        word-spacing: 0.1em;
        letter-spacing: 0.02em;
    }
    
    .ieee-title {
        font-size: 24pt;
        font-weight: bold;
        text-align: center !important;
        margin: 20px 0;
        line-height: 1.3;
    }
    
    .ieee-authors {
        font-size: 10pt;
        text-align: center !important;
        margin: 15px 0;
    }
    
    .ieee-abstract, .ieee-keywords {
        font-size: 9pt;
        font-weight: bold;
        text-align: justify !important;
        text-justify: inter-word !important;
        margin: 12px 0;
        hyphens: auto !important;
    }
    
    .ieee-section-title {
        font-size: 10pt;
        font-weight: bold;
        text-align: center !important;
        text-transform: uppercase;
        margin: 15px 0 5px 0;
    }
    
    .ieee-paragraph {
        font-size: 10pt;
        text-align: justify !important;
        text-justify: inter-word !important;
        text-align-last: justify !important;
        margin: 0 0 12px 0;
        hyphens: auto !important;
        word-spacing: 0.1em;
        letter-spacing: 0.02em;
    }
    
    .ieee-table-caption {
        font-size: 9pt;
        font-weight: bold;
        text-align: center !important;
        margin: 12px 0 6px 0;
        text-transform: uppercase;
    }
    
    .ieee-figure-caption {
        font-size: 9pt;
        font-weight: bold;
        text-align: center !important;
        margin: 6px 0 12px 0;
        text-transform: uppercase;
    }
    
    .ieee-image-container {
        text-align: center !important;
        margin: 12px 0;
        break-inside: avoid;
    }
    
    .ieee-image {
        max-width: 100%;
        height: auto;
        display: block;
        margin: 0 auto;
    }
    
    table {
        border-collapse: collapse;
        margin: 12px auto;
        font-size: 9pt;
    }
    
    th, td {
        border: 1px solid black;
        padding: 4px 6px;
        text-align: left;
        vertical-align: top;
    }
    
    th {
        font-weight: bold;
        text-align: center !important;
        background-color: #f5f5f5;
    }
    
    .ieee-references {
        font-size: 9pt;
        text-align: justify !important;
        text-justify: inter-word !important;
        padding-left: 15px;
        text-indent: -15px;
        margin: 6px 0;
        hyphens: auto !important;
    }
    
    /* Two-column layout */
    .ieee-two-column {
        column-count: 2;
        column-gap: 0.25in;
        column-fill: balance;
    }
    
    .ieee-two-column p {
        text-align: justify !important;
        text-justify: inter-word !important;
        hyphens: auto !important;
        break-inside: avoid-column;
    }
    </style>
    """
    
    # Inject enhanced CSS into HTML
    if '<head>' in html_content:
        html_content = html_content.replace('<head>', f'<head>{enhanced_css}')
    else:
        html_content = f'<html><head>{enhanced_css}</head><body>{html_content}</body></html>'
    
    try:
        import weasyprint
        
        # Generate PDF with WeasyPrint
        html_doc = weasyprint.HTML(string=html_content)
        html_doc.write_pdf(output_path)
        
        return True
        
    except ImportError:
        print("WeasyPrint not available, falling back to basic PDF generation", file=sys.stderr)
        return False
    except Exception as e:
        print(f"Error generating PDF: {e}", file=sys.stderr)
        return False

def main():
    """Main function to generate IEEE document with comprehensive fixes."""
    if len(sys.argv) < 2:
        print("Usage: python ieee_generator_comprehensive_fix.py <input_json>", file=sys.stderr)
        sys.exit(1)
    
    try:
        # Read input JSON
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Create new document
        doc = Document()
        
        # Apply document defaults
        # set_document_defaults(doc)  # Implement this function as needed
        
        # Add title
        if data.get('title'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(sanitize_text(data['title']))
            title_run.bold = True
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(24)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(12)
        
        # Add authors
        if data.get('authors'):
            # Implement author addition logic
            pass
        
        # Add abstract
        if data.get('abstract'):
            abstract_para = doc.add_paragraph()
            abstract_title = abstract_para.add_run("Abstract—")
            abstract_title.bold = True
            abstract_title.font.name = 'Times New Roman'
            abstract_title.font.size = Pt(9)
            
            abstract_content = abstract_para.add_run(sanitize_text(data['abstract']))
            abstract_content.bold = True
            abstract_content.font.name = 'Times New Roman'
            abstract_content.font.size = Pt(9)
            
            apply_perfect_justification(abstract_para)
            abstract_para.paragraph_format.space_after = Pt(6)
        
        # Add keywords
        if data.get('keywords'):
            keywords_para = doc.add_paragraph()
            keywords_title = keywords_para.add_run("Keywords—")
            keywords_title.bold = True
            keywords_title.font.name = 'Times New Roman'
            keywords_title.font.size = Pt(9)
            
            keywords_content = keywords_para.add_run(sanitize_text(data['keywords']))
            keywords_content.bold = True
            keywords_content.font.name = 'Times New Roman'
            keywords_content.font.size = Pt(9)
            
            apply_perfect_justification(keywords_para)
            keywords_para.paragraph_format.space_after = Pt(12)
        
        # Process sections
        if data.get('sections'):
            table_count = 0
            figure_count = 0
            
            for section_idx, section in enumerate(data['sections'], 1):
                # Add section title
                if section.get('title'):
                    section_para = doc.add_paragraph()
                    section_run = section_para.add_run(f"{section_idx}. {sanitize_text(section['title']).upper()}")
                    section_run.bold = True
                    section_run.font.name = 'Times New Roman'
                    section_run.font.size = Pt(10)
                    section_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    section_para.paragraph_format.space_before = Pt(12)
                    section_para.paragraph_format.space_after = Pt(6)
                
                # Process content blocks
                if section.get('contentBlocks'):
                    for block in section['contentBlocks']:
                        block_type = block.get('type', '')
                        
                        if block_type == 'text' and block.get('content'):
                            # Add text paragraph with perfect justification
                            text_para = doc.add_paragraph()
                            text_run = text_para.add_run(sanitize_text(block['content']))
                            text_run.font.name = 'Times New Roman'
                            text_run.font.size = Pt(10)
                            apply_perfect_justification(text_para)
                            text_para.paragraph_format.space_after = Pt(6)
                        
                        elif block_type == 'image' and block.get('data'):
                            # Add image with proper positioning
                            figure_count += 1
                            size = block.get('size', 'medium')
                            size_mapping = {
                                'small': Inches(2.0),
                                'medium': Inches(2.5),
                                'large': Inches(3.3125)
                            }
                            width = size_mapping.get(size, Inches(2.5))
                            
                            caption = block.get('caption', f'Figure {section_idx}.{figure_count}')
                            figure_num = f"{section_idx}.{figure_count}"
                            
                            add_image_with_proper_positioning(doc, block['data'], width, caption, figure_num)
                        
                        elif block_type == 'table':
                            # Add table with fixed caption handling
                            table_count += 1
                            add_ieee_table_fixed(doc, block, section_idx, table_count)
        
        # Save DOCX
        output_docx = sys.argv[1].replace('.json', '_fixed.docx')
        doc.save(output_docx)
        print(f"DOCX saved: {output_docx}")
        
        # Generate HTML for PDF
        html_content = generate_html_from_data(data)  # Implement this function
        
        # Generate PDF with perfect justification
        output_pdf = sys.argv[1].replace('.json', '_fixed.pdf')
        if generate_pdf_with_perfect_justification(html_content, output_pdf):
            print(f"PDF saved: {output_pdf}")
        else:
            print("PDF generation failed")
        
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()