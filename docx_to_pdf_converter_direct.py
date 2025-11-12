#!/usr/bin/env python3
"""
Direct DOCX to PDF Converter with 2-Column IEEE Layout
Pure Python solution using python-docx + reportlab for exact conversion
"""

import sys
import os
import tempfile
import logging
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class IEEE2ColumnDocTemplate(BaseDocTemplate):
    """Custom document template for IEEE 2-column layout"""
    
    def __init__(self, filename, **kwargs):
        BaseDocTemplate.__init__(self, filename, **kwargs)
        
        # Page dimensions
        self.pageWidth = A4[0]  # 595 points
        self.pageHeight = A4[1]  # 842 points
        
        # Margins
        self.leftMargin = 54  # 0.75 inches
        self.rightMargin = 54
        self.topMargin = 72   # 1 inch
        self.bottomMargin = 72
        
        # Column settings
        self.columnGap = 18   # 0.25 inches
        self.columnWidth = (self.pageWidth - self.leftMargin - self.rightMargin - self.columnGap) / 2
        
        # Create frames for 2-column layout
        leftFrame = Frame(
            self.leftMargin, 
            self.bottomMargin,
            self.columnWidth, 
            self.pageHeight - self.topMargin - self.bottomMargin,
            leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
            id='leftColumn'
        )
        
        rightFrame = Frame(
            self.leftMargin + self.columnWidth + self.columnGap,
            self.bottomMargin,
            self.columnWidth,
            self.pageHeight - self.topMargin - self.bottomMargin,
            leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
            id='rightColumn'
        )
        
        # Create page template with 2 columns
        twoColumnTemplate = PageTemplate(
            id='twoColumn',
            frames=[leftFrame, rightFrame]
        )
        
        # Single column template for title and abstract
        singleFrame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.pageWidth - self.leftMargin - self.rightMargin,
            self.pageHeight - self.topMargin - self.bottomMargin,
            leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
            id='singleColumn'
        )
        
        singleColumnTemplate = PageTemplate(
            id='singleColumn',
            frames=[singleFrame]
        )
        
        self.addPageTemplates([singleColumnTemplate, twoColumnTemplate])

def convert_docx_to_pdf_direct(docx_data):
    """
    Convert DOCX to PDF with exact 2-column IEEE formatting
    """
    try:
        logger.info("Starting direct DOCX→PDF conversion with 2-column layout")
        
        # Create temporary file for DOCX input
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_data)
            temp_docx_path = temp_docx.name
        
        try:
            # Load the DOCX document
            logger.info(f"Loading DOCX document from {temp_docx_path}")
            doc = Document(temp_docx_path)
            
            # Create PDF in memory
            pdf_buffer = BytesIO()
            pdf_doc = IEEE2ColumnDocTemplate(pdf_buffer)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles for IEEE format
            title_style = ParagraphStyle(
                'IEEETitle',
                parent=styles['Title'],
                fontSize=14,
                fontName='Helvetica-Bold',
                alignment=TA_CENTER,
                spaceAfter=12,
                spaceBefore=0
            )
            
            author_style = ParagraphStyle(
                'IEEEAuthor',
                parent=styles['Normal'],
                fontSize=12,
                fontName='Helvetica',
                alignment=TA_CENTER,
                spaceAfter=18,
                spaceBefore=6
            )
            
            abstract_style = ParagraphStyle(
                'IEEEAbstract',
                parent=styles['Normal'],
                fontSize=9,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                spaceAfter=12,
                spaceBefore=6,
                leftIndent=0,
                rightIndent=0
            )
            
            section_heading_style = ParagraphStyle(
                'IEEESectionHeading',
                parent=styles['Heading2'],
                fontSize=10,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                spaceAfter=6,
                spaceBefore=12,
                textTransform='uppercase'
            )
            
            body_style = ParagraphStyle(
                'IEEEBody',
                parent=styles['Normal'],
                fontSize=9,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                spaceAfter=6,
                spaceBefore=0,
                leftIndent=0,
                rightIndent=0
            )
            
            story = []
            
            # Process document content
            title_found = False
            abstract_content = ""
            in_abstract = False
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect title (first significant paragraph)
                if not title_found and len(text) > 5:
                    story.append(Paragraph(text, title_style))
                    title_found = True
                    continue
                
                # Detect abstract
                if text.lower().startswith('abstract'):
                    in_abstract = True
                    # Extract abstract text after "Abstract" keyword
                    abstract_text = text.replace('Abstract', '').replace('abstract', '').strip()
                    if abstract_text.startswith('—') or abstract_text.startswith('-'):
                        abstract_text = abstract_text[1:].strip()
                    if abstract_text:
                        abstract_content = abstract_text
                    continue
                
                if in_abstract:
                    # Check if we've reached the end of abstract (keywords, section heading, etc.)
                    if (text.lower().startswith('keywords') or 
                        text.lower().startswith('index terms') or
                        text.startswith('I.') or text.startswith('1.')):
                        in_abstract = False
                        # Add the abstract to story
                        if abstract_content:
                            story.append(Paragraph(f"<b>Abstract—</b>{abstract_content}", abstract_style))
                            story.append(Spacer(1, 12))
                        # Process this line as normal content
                    else:
                        abstract_content += " " + text
                        continue
                
                # Detect section headings (Roman numerals or numbers)
                if (text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.')) or
                    text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
                    
                    # Switch to 2-column layout if not already
                    if len(story) > 0 and not any('NextPageTemplate' in str(item) for item in story[-3:]):
                        story.append(PageBreak())
                        story.append(Paragraph('<para><nextPageTemplate name="twoColumn"/></para>', styles['Normal']))
                    
                    story.append(Paragraph(text, section_heading_style))
                else:
                    # Regular body text
                    story.append(Paragraph(text, body_style))
            
            # Add final abstract if we were still in it
            if in_abstract and abstract_content:
                story.insert(-len([item for item in story if 'NextPageTemplate' not in str(item)]), 
                           Paragraph(f"<b>Abstract—</b>{abstract_content}", abstract_style))
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    # Create reportlab table
                    t = Table(table_data, colWidths=[pdf_doc.columnWidth/len(table_data[0])] * len(table_data[0]))
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP')
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            pdf_data = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            logger.info(f"Direct DOCX→PDF conversion successful, size: {len(pdf_data)} bytes")
            return pdf_data
                    
        finally:
            # Clean up DOCX temp file
            try:
                if os.path.exists(temp_docx_path):
                    os.unlink(temp_docx_path)
            except:
                pass
            
    except Exception as e:
        logger.error(f"Direct DOCX→PDF conversion failed: {e}")
        raise Exception(f"DOCX→PDF conversion failed: {str(e)}")

# Backward compatibility functions
def convert_docx_to_pdf(docx_data):
    """Main conversion function"""
    return convert_docx_to_pdf_direct(docx_data)

def convert_docx_file_to_pdf(input_docx_path, output_pdf_path):
    """Convert DOCX file to PDF file"""
    try:
        with open(input_docx_path, 'rb') as f:
            docx_data = f.read()
        
        pdf_data = convert_docx_to_pdf_direct(docx_data)
        
        with open(output_pdf_path, 'wb') as f:
            f.write(pdf_data)
        
        logger.info(f"File conversion successful: {input_docx_path} -> {output_pdf_path}")
        
    except Exception as e:
        logger.error(f"File conversion failed: {e}")
        raise

if __name__ == "__main__":
    # Simple test
    print("Direct DOCX→PDF converter ready")