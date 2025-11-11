#!/usr/bin/env python3
"""
DOCX to PDF Converter
Converts Word documents to PDF format using ReportLab (Vercel-compatible)
"""

import sys
import os
import tempfile
import logging
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib import colors
from PIL import Image

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def extract_docx_content(docx_path):
    """
    Extract content from DOCX file for PDF conversion
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        List of content elements (paragraphs, tables, images)
    """
    try:
        doc = Document(docx_path)
        content = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        content.append({
                            'type': 'paragraph',
                            'text': para.text,
                            'alignment': para.alignment,
                            'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic} for run in para.runs]
                        })
                        break
                        
            elif element.tag.endswith('tbl'):  # Table
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text)
                            table_data.append(row_data)
                        content.append({
                            'type': 'table',
                            'data': table_data
                        })
                        break
        
        return content
        
    except Exception as e:
        logger.error(f"Error extracting DOCX content: {e}")
        raise

def convert_alignment(docx_alignment):
    """Convert DOCX alignment to ReportLab alignment"""
    if docx_alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return TA_CENTER
    elif docx_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return TA_RIGHT
    elif docx_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return TA_JUSTIFY
    else:
        return TA_LEFT

def create_pdf_from_content(content, output_path):
    """
    Create PDF from extracted DOCX content using ReportLab
    
    Args:
        content: List of content elements
        output_path: Path for output PDF
    """
    try:
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        story = []
        
        for item in content:
            if item['type'] == 'paragraph':
                if item['text'].strip():  # Only add non-empty paragraphs
                    # Create paragraph style based on formatting
                    style = styles['Normal']
                    if item.get('alignment'):
                        alignment = convert_alignment(item['alignment'])
                        style = ParagraphStyle(
                            'CustomStyle',
                            parent=styles['Normal'],
                            alignment=alignment
                        )
                    
                    # Handle formatted text
                    formatted_text = ""
                    for run in item.get('runs', []):
                        text = run['text']
                        if run.get('bold'):
                            text = f"<b>{text}</b>"
                        if run.get('italic'):
                            text = f"<i>{text}</i>"
                        formatted_text += text
                    
                    if not formatted_text.strip():
                        formatted_text = item['text']
                    
                    para = Paragraph(formatted_text, style)
                    story.append(para)
                    story.append(Spacer(1, 6))
                    
            elif item['type'] == 'table':
                if item['data']:
                    # Create table
                    table = Table(item['data'])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        logger.info(f"PDF created successfully: {output_path}")
        
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        raise

def convert_docx_file_to_pdf(input_docx_path, output_pdf_path):
    """
    Convert DOCX file to PDF file using ReportLab
    
    Args:
        input_docx_path: Path to input DOCX file
        output_pdf_path: Path to output PDF file
    """
    try:
        logger.info(f"Starting DOCX to PDF conversion: {input_docx_path} -> {output_pdf_path}")
        
        # Check if input file exists
        if not os.path.exists(input_docx_path):
            raise FileNotFoundError(f"Input DOCX file not found: {input_docx_path}")
        
        # Get file size for logging
        file_size = os.path.getsize(input_docx_path)
        logger.info(f"Input DOCX file size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("Input DOCX file is empty")
        
        # Extract content from DOCX
        content = extract_docx_content(input_docx_path)
        
        # Create PDF from content
        create_pdf_from_content(content, output_pdf_path)
        
        # Check if output file was created
        if not os.path.exists(output_pdf_path):
            raise RuntimeError("PDF file was not created")
        
        # Check output file size
        output_size = os.path.getsize(output_pdf_path)
        logger.info(f"PDF conversion successful, output size: {output_size} bytes")
        
        if output_size == 0:
            raise RuntimeError("Generated PDF file is empty")
            
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {e}")
        raise

def convert_docx_to_pdf(docx_data):
    """
    Convert DOCX data to PDF (legacy function for backward compatibility)
    
    Args:
        docx_data: Binary data of the DOCX file
    
    Returns:
        Binary PDF data
    """
    try:
        logger.info("Starting DOCX to PDF conversion from binary data")
        
        # Create temporary files for input and output
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            # Write DOCX data to temporary file
            temp_docx.write(docx_data)
            temp_docx_path = temp_docx.name
            logger.info(f"Temporary DOCX file created: {temp_docx_path}")
        
        # Create temporary PDF file path
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            # Convert using the file-based function
            convert_docx_file_to_pdf(temp_docx_path, temp_pdf_path)
            
            # Read the generated PDF
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                logger.info(f"PDF data read, size: {len(pdf_data)} bytes")
            
            return pdf_data
            
        finally:
            # Clean up temporary files
            try:
                if os.path.exists(temp_docx_path):
                    os.unlink(temp_docx_path)
                    logger.info(f"Cleaned up temporary DOCX file: {temp_docx_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up DOCX file: {e}")
                
            try:
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
                    logger.info(f"Cleaned up temporary PDF file: {temp_pdf_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up PDF file: {e}")
                
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {e}")
        raise

def main():
    """
    Main function to handle conversion
    Can work in two modes:
    1. File paths as command line arguments: python converter.py input.docx output.pdf
    2. Legacy stdin mode (for backward compatibility)
    """
    try:
        logger.info("DOCX to PDF converter started (ReportLab-based)")
        
        # Check if file paths are provided as arguments
        if len(sys.argv) == 3:
            # File path mode
            input_docx_path = sys.argv[1]
            output_pdf_path = sys.argv[2]
            
            logger.info(f"File path mode: {input_docx_path} -> {output_pdf_path}")
            convert_docx_file_to_pdf(input_docx_path, output_pdf_path)
            logger.info("DOCX to PDF conversion completed successfully")
            
        elif len(sys.argv) == 1:
            # Legacy stdin mode
            logger.info("Legacy stdin mode")
            
            # Read DOCX data from stdin
            docx_data = sys.stdin.buffer.read()
            logger.info(f"Read {len(docx_data)} bytes of DOCX data from stdin")
            
            if len(docx_data) == 0:
                logger.error("No DOCX data received from stdin")
                sys.stderr.write("Error: No DOCX data received\n")
                sys.exit(1)
            
            # Convert to PDF
            pdf_data = convert_docx_to_pdf(docx_data)
            
            if len(pdf_data) == 0:
                logger.error("PDF conversion resulted in empty file")
                sys.stderr.write("Error: PDF conversion resulted in empty file\n")
                sys.exit(1)
            
            # Write PDF data to stdout
            sys.stdout.buffer.write(pdf_data)
            sys.stdout.buffer.flush()
            
            logger.info("DOCX to PDF conversion completed successfully")
            
        else:
            logger.error("Invalid number of arguments")
            sys.stderr.write("Usage: python converter.py [input.docx output.pdf] or pipe DOCX data to stdin\n")
            sys.exit(1)
        
    except Exception as e:
        logger.error(f"Conversion error: {e}")
        sys.stderr.write(f"Error: {e}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()