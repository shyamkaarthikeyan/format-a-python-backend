#!/usr/bin/env python3
"""
Minimal DOCX test to isolate the 500 error issue
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def create_minimal_docx(document_data):
    """Create a minimal DOCX document to test basic functionality"""
    try:
        print("Creating minimal DOCX document...", file=sys.stderr)
        
        # Create new document
        doc = Document()
        
        # Set basic margins
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        # Add title
        if document_data.get('title'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(str(document_data['title']))
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(24)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add authors
        if document_data.get('authors'):
            for author in document_data['authors']:
                if author.get('name'):
                    author_para = doc.add_paragraph()
                    author_run = author_para.add_run(str(author['name']))
                    author_run.font.name = 'Times New Roman'
                    author_run.font.size = Pt(10)
                    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add abstract
        if document_data.get('abstract'):
            abstract_para = doc.add_paragraph()
            abstract_title = abstract_para.add_run("Abstract—")
            abstract_title.font.name = 'Times New Roman'
            abstract_title.font.size = Pt(9)
            abstract_title.bold = True
            
            abstract_content = abstract_para.add_run(str(document_data['abstract']))
            abstract_content.font.name = 'Times New Roman'
            abstract_content.font.size = Pt(9)
            abstract_content.bold = True
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        print("Minimal DOCX created successfully", file=sys.stderr)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating minimal DOCX: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        raise

if __name__ == "__main__":
    # Test the function
    test_data = {
        "title": "Test Document",
        "authors": [{"name": "Test Author"}],
        "abstract": "Test abstract content"
    }
    
    try:
        docx_bytes = create_minimal_docx(test_data)
        print(f"Success! Generated {len(docx_bytes)} bytes")
    except Exception as e:
        print(f"Failed: {e}")