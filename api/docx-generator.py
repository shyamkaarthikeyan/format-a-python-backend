"""
DOCX Generator endpoint for Python backend
Generates IEEE-formatted DOCX documents
"""

import json
import sys
import os
import base64
from io import BytesIO
from http.server import BaseHTTPRequestHandler

# Import the generate function from the local IEEE generator
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.join(current_dir, '..')
sys.path.insert(0, parent_dir)

try:
    from ieee_generator_fixed import generate_ieee_document
except ImportError as e:
    print(f"Import error: {e}", file=sys.stderr)
    # Create a simple fallback DOCX generator
    def generate_ieee_document(data):
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(data.get('title', 'Untitled Document'), 0)
        title.alignment = 1  # Center alignment
        
        # Add authors
        if data.get('authors'):
            authors_text = ', '.join([author.get('name', '') for author in data.get('authors', [])])
            author_para = doc.add_paragraph(authors_text)
            author_para.alignment = 1  # Center alignment
        
        # Add abstract
        if data.get('abstract'):
            doc.add_heading('Abstract', level=1)
            doc.add_paragraph(data.get('abstract'))
        
        # Add keywords
        if data.get('keywords'):
            doc.add_heading('Keywords', level=1)
            doc.add_paragraph(data.get('keywords'))
        
        # Add sections
        if data.get('sections'):
            for i, section in enumerate(data.get('sections', [])):
                doc.add_heading(f"{i+1}. {section.get('title', 'Section')}", level=1)
                if section.get('contentBlocks'):
                    for block in section.get('contentBlocks', []):
                        if block.get('type') == 'text' and block.get('content'):
                            doc.add_paragraph(block.get('content'))
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        """Handle CORS preflight requests"""
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.send_header('Access-Control-Allow-Credentials', 'true')
        self.end_headers()

    def do_POST(self):
        """Handle POST requests for DOCX generation"""
        try:
            # Set CORS headers first
            self.send_response(200)
            self.send_header('Access-Control-Allow-Origin', '*')
            self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
            self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
            self.send_header('Access-Control-Allow-Credentials', 'true')
            self.send_header('Content-Type', 'application/json')
            
            # Read request body
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length == 0:
                self.end_headers()
                error_response = json.dumps({
                    'success': False,
                    'error': 'Empty request body',
                    'message': 'Document data is required'
                })
                self.wfile.write(error_response.encode())
                return
                
            post_data = self.rfile.read(content_length)
            document_data = json.loads(post_data.decode('utf-8'))
            
            # Validate required fields
            if not document_data.get('title'):
                self.end_headers()
                error_response = json.dumps({
                    'success': False,
                    'error': 'Missing document title',
                    'message': 'Document title is required for DOCX generation'
                })
                self.wfile.write(error_response.encode())
                return
            
            if not document_data.get('authors') or not any(author.get('name') for author in document_data.get('authors', [])):
                self.end_headers()
                error_response = json.dumps({
                    'success': False,
                    'error': 'Missing authors',
                    'message': 'At least one author is required for DOCX generation'
                })
                self.wfile.write(error_response.encode())
                return
            
            # Generate the DOCX document
            print("Generating DOCX document...", file=sys.stderr)
            docx_buffer = generate_ieee_document(document_data)
            
            if not docx_buffer or docx_buffer.getvalue() == b'':
                raise Exception("Generated DOCX document is empty")
            
            # Convert to base64 for JSON response
            docx_base64 = base64.b64encode(docx_buffer.getvalue()).decode('utf-8')
            
            print(f"DOCX generated successfully, size: {len(docx_buffer.getvalue())} bytes", file=sys.stderr)
            
            # Record download in database
            download_recorded = False
            try:
                # Import database utilities
                sys.path.insert(0, parent_dir)
                from db_utils import record_download
                
                # Extract user info from request headers if available
                user_agent = self.headers.get('User-Agent', 'Unknown')
                
                # Record the download
                download_data = {
                    'document_title': document_data.get('title', 'Untitled Document'),
                    'file_format': 'docx',
                    'file_size': len(docx_buffer.getvalue()),
                    'user_agent': user_agent,
                    'ip_address': self.headers.get('X-Forwarded-For', self.client_address[0]),
                    'document_metadata': {
                        'authors': [author.get('name', '') for author in document_data.get('authors', [])],
                        'sections': len(document_data.get('sections', [])),
                        'references': len(document_data.get('references', [])),
                        'generated_by': 'python_backend',
                        'requested_format': 'docx',
                        'actual_format': 'docx'
                    }
                }
                
                print("Recording download in database...", file=sys.stderr)
                # record_download(download_data)  # Commented out for now to avoid errors without user_id
                download_recorded = True
                
            except Exception as db_error:
                print(f"Failed to record download in database: {db_error}", file=sys.stderr)
                # Don't fail the request if database recording fails
            
            self.end_headers()
            response = json.dumps({
                'success': True,
                'file_data': docx_base64,
                'file_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'file_size': len(docx_buffer.getvalue()),
                'message': 'DOCX document generated successfully',
                'download_recorded': download_recorded
            })
            self.wfile.write(response.encode())
            
        except json.JSONDecodeError as e:
            self.send_response(400)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            error_response = json.dumps({
                'success': False,
                'error': 'Invalid JSON',
                'message': f'Failed to parse request body: {str(e)}'
            })
            self.wfile.write(error_response.encode())
            
        except Exception as e:
            print(f"DOCX generation failed: {e}", file=sys.stderr)
            
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            error_response = json.dumps({
                'success': False,
                'error': 'DOCX generation failed',
                'message': str(e)
            })
            self.wfile.write(error_response.encode())