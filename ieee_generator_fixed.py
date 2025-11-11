"""
IEEE Document Generator - EXACT copy from test.py
"""

import argparse
import base64
import json
import os
import re
import subprocess
import sys
import tempfile
import unicodedata
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def sanitize_text(text):
    """Sanitize text to remove invalid Unicode characters and surrogates."""
    if not text:
        return ""

    # Convert to string if not already
    text = str(text)

    # Remove surrogate characters and other problematic Unicode
    text = text.encode("utf-8", "ignore").decode("utf-8")

    # Normalize Unicode characters
    text = unicodedata.normalize("NFKD", text)

    # Remove any remaining control characters except newlines and tabs
    text = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]", "", text)

    return text


def add_image_with_proper_layout(doc, image_data, width, caption_text="", figure_number=""):
    """Add image with proper layout optimized for 2-column IEEE format."""
    try:
        # Decode base64 image data
        if "," in image_data:
            image_data = image_data.split(",")[1]
        image_bytes = base64.b64decode(image_data)
        image_stream = BytesIO(image_bytes)
        
        # Ensure width fits within 2-column layout constraints
        max_column_width = Inches(3.0)  # Safe width for 2-column layout
        if width > max_column_width:
            width = max_column_width
        
        # Add spacing before image
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(12)
        spacing_para.paragraph_format.space_after = Pt(6)
        
        # Create paragraph for image with left alignment for better 2-column compatibility
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment works better in narrow columns
        img_para.paragraph_format.space_before = Pt(6)
        img_para.paragraph_format.space_after = Pt(6)
        
        # Add image to paragraph with comprehensive visibility settings
        run = img_para.add_run()
        picture = run.add_picture(image_stream, width=width)
        
        # CRITICAL: Set image positioning properties for maximum visibility in 2-column layout
        inline = picture._inline
        
        # Force image to be visible and properly positioned
        docPr = inline.docPr
        docPr.set('name', f'Image_{figure_number}' if figure_number else 'Image')
        docPr.set('descr', caption_text if caption_text else 'Figure')
        
        # Set graphic properties for maximum visibility
        graphic = inline.graphic
        graphicData = graphic.graphicData
        pic = graphicData.pic
        
        # Ensure image has proper display properties
        spPr = pic.spPr
        if spPr is not None:
            # Add transform properties for proper positioning
            xfrm = spPr.xfrm
            if xfrm is not None:
                # Ensure image is not clipped or offset incorrectly
                off = xfrm.off
                if off is not None:
                    off.set('x', '0')
                    off.set('y', '0')
                
                # Set proper extent for visibility
                ext = xfrm.ext
                if ext is not None:
                    # Ensure image extent is properly set
                    cx = int(width.emu)  # Convert to EMUs
                    ext.set('cx', str(cx))
        
        # Scale down if too tall to prevent page overflow
        if picture.height > IEEE_CONFIG["max_figure_height"]:
            scale_factor = IEEE_CONFIG["max_figure_height"] / picture.height
            run.clear()
            image_stream.seek(0)  # Reset stream position
            picture = run.add_picture(
                image_stream,
                width=width * scale_factor,
                height=IEEE_CONFIG["max_figure_height"],
            )
            
            # Reapply visibility settings after scaling
            inline = picture._inline
            docPr = inline.docPr
            docPr.set('name', f'Image_{figure_number}_scaled' if figure_number else 'Image_scaled')
            docPr.set('descr', f'{caption_text} (scaled)' if caption_text else 'Figure (scaled)')
        
        # Add caption if provided
        if caption_text:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment for 2-column compatibility
            caption_para.paragraph_format.space_before = Pt(3)
            caption_para.paragraph_format.space_after = Pt(12)
            
            caption_run = caption_para.add_run()
            if figure_number:
                caption_run.text = f"{figure_number}: {caption_text}"
            else:
                caption_run.text = caption_text
            caption_run.font.size = IEEE_CONFIG["font_size_caption"]
            caption_run.font.name = IEEE_CONFIG["font_name"]
            caption_run.italic = True
            caption_run.bold = True
        
        # Add spacing after image
        spacing_para_after = doc.add_paragraph()
        spacing_para_after.paragraph_format.space_before = Pt(6)
        spacing_para_after.paragraph_format.space_after = Pt(12)
        
        return True
        
    except Exception as e:
        print(f"Error adding image with proper layout: {e}", file=sys.stderr)
        return False


# IEEE EXACT LATEX PDF FORMATTING - LOW-LEVEL OPENXML SPECIFICATIONS
IEEE_CONFIG = {
    "font_name": "Times New Roman",
    # Font sizes (exact IEEE LaTeX specifications)
    "font_size_title": Pt(24),  # Title: 24pt bold centered
    "font_size_author_name": Pt(10),  # Author names: 10pt bold
    "font_size_author_affil": Pt(10),  # Author affiliations: 10pt italic
    "font_size_author_email": Pt(9),  # Author emails: 9pt
    "font_size_body": Pt(10),  # Body text: 10pt
    "font_size_abstract": Pt(9),  # Abstract/Keywords: 9pt bold
    "font_size_caption": Pt(9),  # Captions: 9pt italic
    "font_size_reference": Pt(9),  # References: 9pt
    # Page margins (exact IEEE LaTeX: 0.75" all sides = 1080 twips)
    "margin_twips": 1080,
    # Two-column layout (exact IEEE LaTeX specifications)
    "column_count": 2,
    "column_width_twips": 4770,  # 3.3125" per column
    "column_gap_twips": 360,  # 0.25" gap between columns
    "column_indent": Pt(0),  # No indent for column text
    # Line spacing (exact IEEE LaTeX: 12pt = 240 twips)
    "line_spacing_twips": 240,
    "line_spacing": Pt(12),  # For backward compatibility
    # Paragraph spacing (exact IEEE LaTeX specifications)
    "spacing_title_after": 240,  # 12pt after title
    "spacing_abstract_after": 120,  # 6pt after abstract
    "spacing_keywords_after": 240,  # 12pt after keywords
    "spacing_section_before": 240,  # 12pt before section headings
    "spacing_section_after": 0,  # 0pt after section headings
    # Figure specifications
    "figure_max_width_twips": 4770,  # Max 3.3125" width (column width)
    "figure_spacing": 120,  # 6pt before/after figures
    "figure_sizes": {
        "Very Small": Inches(1.5),  # 1.5" width
        "Small": Inches(2.0),  # 2.0" width
        "Medium": Inches(2.5),  # 2.5" width
        "Large": Inches(3.3125),  # Full column width
    },
    "max_figure_height": Inches(4.0),  # Max figure height
    # Reference specifications
    "reference_hanging_indent": 360,  # 0.25" hanging indent
}


def set_document_defaults(doc):
    """Set document-wide defaults using EXACT IEEE LaTeX PDF specifications via OpenXML."""

    # 1. SET PAGE MARGINS - EXACT IEEE LaTeX: 0.75" all sides (1080 twips)
    for section in doc.sections:
        sectPr = section._sectPr
        pgMar = (
            sectPr.xpath("./w:pgMar")[0]
            if sectPr.xpath("./w:pgMar")
            else OxmlElement("w:pgMar")
        )
        pgMar.set(qn("w:left"), "1080")
        pgMar.set(qn("w:right"), "1080")
        pgMar.set(qn("w:top"), "1080")
        pgMar.set(qn("w:bottom"), "1080")
        if not sectPr.xpath("./w:pgMar"):
            sectPr.append(pgMar)

    # 2. SET HYPHENATION & COMPATIBILITY - EXACT IEEE LaTeX specifications
    settings = doc.settings
    settings_element = settings.element

    # Clear existing settings first
    for elem in settings_element.xpath(
        "./w:autoHyphenation | ./w:hyphenationZone | ./w:consecutiveHyphenLimit | ./w:doNotHyphenateCaps | ./w:compat"
    ):
        settings_element.remove(elem)

    # Add hyphenation settings
    autoHyphenation = OxmlElement("w:autoHyphenation")
    autoHyphenation.set(qn("w:val"), "1")
    settings_element.append(autoHyphenation)

    hyphenationZone = OxmlElement("w:hyphenationZone")
    hyphenationZone.set(qn("w:val"), "360")  # 0.25"
    settings_element.append(hyphenationZone)

    consecutiveHyphenLimit = OxmlElement("w:consecutiveHyphenLimit")
    consecutiveHyphenLimit.set(qn("w:val"), "2")
    settings_element.append(consecutiveHyphenLimit)

    doNotHyphenateCaps = OxmlElement("w:doNotHyphenateCaps")
    doNotHyphenateCaps.set(qn("w:val"), "1")
    settings_element.append(doNotHyphenateCaps)

    # Add compatibility settings
    compat = OxmlElement("w:compat")

    usePrinterMetrics = OxmlElement("w:usePrinterMetrics")
    usePrinterMetrics.set(qn("w:val"), "1")
    compat.append(usePrinterMetrics)

    doNotExpandShiftReturn = OxmlElement("w:doNotExpandShiftReturn")
    doNotExpandShiftReturn.set(qn("w:val"), "1")
    compat.append(doNotExpandShiftReturn)

    settings_element.append(compat)

    # 3. MODIFY NORMAL STYLE - EXACT IEEE LaTeX specifications via OpenXML
    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        # Apply via OpenXML for exact control
        style_element = normal.element
        pPr = (
            style_element.xpath("./w:pPr")[0]
            if style_element.xpath("./w:pPr")
            else OxmlElement("w:pPr")
        )

        # Clear existing spacing/alignment
        for elem in pPr.xpath("./w:spacing | ./w:jc | ./w:ind"):
            pPr.remove(elem)

        # EXACT 12pt line spacing (240 twips)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

        # Full justification
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

        if not style_element.xpath("./w:pPr"):
            style_element.append(pPr)

        # Font settings
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10)

    # Modify Heading 1 style - IEEE SECTION HEADINGS (BOLD, CENTERED, UPPERCASE)
    if "Heading 1" in styles:
        heading1 = styles["Heading 1"]
        style_element = heading1.element
        pPr = (
            style_element.xpath("./w:pPr")[0]
            if style_element.xpath("./w:pPr")
            else OxmlElement("w:pPr")
        )

        # Clear existing formatting
        for elem in pPr.xpath("./w:spacing | ./w:jc"):
            pPr.remove(elem)

        # Section heading spacing: 12pt before, 0pt after
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "240")  # 12pt
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

        # Center alignment
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)

        if not style_element.xpath("./w:pPr"):
            style_element.append(pPr)

        heading1.font.name = "Times New Roman"
        heading1.font.size = Pt(10)
        heading1.font.bold = True

    # Modify Heading 2 style for subsections - IEEE SUBSECTION HEADINGS (BOLD, LEFT)
    if "Heading 2" in styles:
        heading2 = styles["Heading 2"]
        style_element = heading2.element
        pPr = (
            style_element.xpath("./w:pPr")[0]
            if style_element.xpath("./w:pPr")
            else OxmlElement("w:pPr")
        )

        # Clear existing formatting
        for elem in pPr.xpath("./w:spacing | ./w:jc"):
            pPr.remove(elem)

        # Subsection spacing: 6pt before, 0pt after
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")  # 6pt
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

        # Left alignment
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)

        if not style_element.xpath("./w:pPr"):
            style_element.append(pPr)

        heading2.font.name = "Times New Roman"
        heading2.font.size = Pt(10)
        heading2.font.bold = True


def add_title(doc, title):
    """Add the paper title - EXACT MATCH TO PDF: 24pt bold centered Times New Roman."""
    para = doc.add_paragraph()
    run = para.add_run(sanitize_text(title))
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(24)  # 24pt exactly like PDF

    # Apply exact OpenXML formatting
    pPr = para._element.get_or_add_pPr()

    # Clear existing formatting
    for elem in pPr.xpath("./w:spacing | ./w:jc"):
        pPr.remove(elem)

    # Center alignment
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)

    # Title spacing: 0pt before, 12pt after
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "240")  # 12pt
    pPr.append(spacing)


def add_authors(doc, authors):
    """Add authors with proper IEEE formatting - 3 authors per row, multiple rows if needed."""
    if not authors:
        return

    if len(authors) == 0:
        return

    # IEEE format: maximum 3 authors per row, create additional rows for more authors
    authors_per_row = 3
    total_authors = len(authors)

    # Process authors in groups of 3
    for row_start in range(0, total_authors, authors_per_row):
        row_end = min(row_start + authors_per_row, total_authors)
        row_authors = authors[row_start:row_end]
        num_cols = len(row_authors)

        # Create table for this row of authors
        table = doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.allow_autofit = True

        # Set table width to full page width
        table.style = "Table Grid"
        table.style.font.name = IEEE_CONFIG["font_name"]
        table.style.font.size = IEEE_CONFIG["font_size_body"]

        # Remove table borders for clean IEEE look
        for table_row in table.rows:
            for cell in table_row.cells:
                # Remove all borders
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in [
                    "top",
                    "left",
                    "bottom",
                    "right",
                    "insideH",
                    "insideV",
                ]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "nil")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # Process each author in this row
        for col_idx, author in enumerate(row_authors):
            if not author.get("name"):
                continue

            cell = table.cell(0, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Set equal column widths - always use 3 column layout for consistency
            cell.width = Inches(
                6.5 / 3
            )  # Always divide by 3 for consistent column width

            # Add proper cell margins for IEEE spacing
            cell_element = cell._element
            cell_properties = cell_element.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")

            # Set cell margins for proper IEEE author block spacing
            for side in ["left", "right", "top", "bottom"]:
                margin = OxmlElement(f"w:{side}")
                margin.set(qn("w:w"), "72")  # 0.05 inch margins
                margin.set(qn("w:type"), "dxa")
                margins.append(margin)
            cell_properties.append(margins)

            # Clear any existing content
            cell._element.clear_content()

            # Author name - bold, centered (IEEE standard)
            name_para = cell.add_paragraph()
            name_run = name_para.add_run(sanitize_text(author["name"]))
            name_run.bold = True  # IEEE standard: author names are bold
            name_run.font.name = IEEE_CONFIG["font_name"]
            name_run.font.size = IEEE_CONFIG["font_size_body"]
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.paragraph_format.space_before = Pt(0)
            name_para.paragraph_format.space_after = Pt(3)

            # Handle individual fields - department, organization, city, etc.
            fields = [
                ("department", "Department"),
                ("organization", "Organization"),
                ("university", "University"),
                ("institution", "Institution"),
                ("city", "City"),
                ("state", "State"),
                ("country", "Country"),
            ]

            # Add each field as a separate line if present
            for field_key, field_name in fields:
                if author.get(field_key):
                    field_para = cell.add_paragraph()
                    field_run = field_para.add_run(sanitize_text(author[field_key]))
                    field_run.italic = True  # IEEE standard: affiliations are italic
                    field_run.font.name = IEEE_CONFIG["font_name"]
                    field_run.font.size = IEEE_CONFIG["font_size_body"]
                    field_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    field_para.paragraph_format.space_before = Pt(0)
                    field_para.paragraph_format.space_after = Pt(2)

            # Process affiliation string if present (for backward compatibility)
            if author.get("affiliation"):
                affiliation_text = author["affiliation"]
                if isinstance(affiliation_text, str):
                    affiliation_lines = affiliation_text.strip().split("\n")
                    for line in affiliation_lines:
                        line = line.strip()
                        if line and not line.lower().startswith(
                            "email"
                        ):  # Skip email lines here
                            affil_para = cell.add_paragraph()
                            affil_run = affil_para.add_run(sanitize_text(line))
                            affil_run.italic = (
                                True  # IEEE standard: affiliations are italic
                            )
                            affil_run.font.name = IEEE_CONFIG["font_name"]
                            affil_run.font.size = IEEE_CONFIG["font_size_body"]
                            affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            affil_para.paragraph_format.space_before = Pt(0)
                            affil_para.paragraph_format.space_after = Pt(2)

            # Handle email field
            email = author.get("email", "")
            if email:
                email_para = cell.add_paragraph()
                email_run = email_para.add_run(sanitize_text(email))
                email_run.font.name = IEEE_CONFIG["font_name"]
                email_run.font.size = Pt(9)  # Slightly smaller for email
                email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                email_para.paragraph_format.space_before = Pt(2)
                email_para.paragraph_format.space_after = Pt(0)

            # Add custom fields if any
            for custom_field in author.get("custom_fields", []):
                if custom_field.get("value"):
                    custom_para = cell.add_paragraph()
                    custom_run = custom_para.add_run(
                        sanitize_text(custom_field["value"])
                    )
                    custom_run.italic = True
                    custom_run.font.name = IEEE_CONFIG["font_name"]
                    custom_run.font.size = IEEE_CONFIG["font_size_body"]
                    custom_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    custom_para.paragraph_format.space_before = Pt(0)
                    custom_para.paragraph_format.space_after = Pt(2)

        # Add spacing between author rows (but not after the last row)
        if row_end < total_authors:
            spacing_para = doc.add_paragraph()
            spacing_para.paragraph_format.space_after = Pt(
                8
            )  # Space between author rows

    # Add proper spacing after all authors
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(12)  # IEEE standard spacing


def add_abstract(doc, abstract):
    """Add the abstract section - EXACT VISUAL MATCH TO PDF: 9pt bold with Abstract— prefix."""
    if abstract:
        # Add abstract with bold title and content in same paragraph - EXACTLY like PDF
        para = doc.add_paragraph()

        # Bold "Abstract—" title (EXACT MATCH TO PDF)
        title_run = para.add_run("Abstract—")
        title_run.bold = True  # PDF uses BOLD for Abstract—
        title_run.italic = False
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(9)  # 9pt like PDF

        # Add abstract content immediately after on SAME LINE (BOLD like PDF)
        content_run = para.add_run(sanitize_text(abstract))
        content_run.bold = True  # PDF uses BOLD for abstract content
        content_run.italic = False
        content_run.font.name = "Times New Roman"
        content_run.font.size = Pt(9)  # 9pt like PDF

        # Apply exact OpenXML formatting
        pPr = para._element.get_or_add_pPr()

        # Clear existing formatting
        for elem in pPr.xpath("./w:spacing | ./w:jc"):
            pPr.remove(elem)

        # Full justification
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

        # Abstract spacing: 0pt before, 6pt after
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "120")  # 6pt
        pPr.append(spacing)


def add_keywords(doc, keywords):
    """Add the keywords section - EXACT IEEE LaTeX PDF formatting with OpenXML."""
    if keywords:
        # Add keywords with bold title and content in same paragraph
        para = doc.add_paragraph()

        # Bold "Keywords—" title (EXACT MATCH TO PDF)
        title_run = para.add_run("Keywords—")
        title_run.bold = True  # PDF uses BOLD for Keywords—
        title_run.italic = False
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(9)  # 9pt like PDF

        # Add keywords content immediately after on SAME LINE (BOLD like PDF)
        content_run = para.add_run(sanitize_text(keywords))
        content_run.bold = True  # PDF uses BOLD for keywords content
        content_run.italic = False
        content_run.font.name = "Times New Roman"
        content_run.font.size = Pt(9)  # 9pt like PDF

        # Apply exact OpenXML formatting
        pPr = para._element.get_or_add_pPr()

        # Clear existing formatting
        for elem in pPr.xpath("./w:spacing | ./w:jc"):
            pPr.remove(elem)

        # Full justification
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

        # Keywords spacing: 0pt before, 12pt after
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "240")  # 12pt
        pPr.append(spacing)


def apply_ieee_latex_formatting(
    para, spacing_before=0, spacing_after=0, line_spacing=240
):
    """Apply PERFECT IEEE LaTeX PDF formatting with enhanced justification controls."""
    pPr = para._element.get_or_add_pPr()

    # Clear existing formatting first
    for elem in pPr.xpath(
        "./w:spacing | ./w:jc | ./w:adjustRightInd | ./w:snapToGrid | ./w:textAlignment | ./w:suppressAutoHyphens"
    ):
        pPr.remove(elem)

    # 1. FULL JUSTIFICATION = CONSISTENT WITH PDF OUTPUT
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")  # Use 'both' for consistent Word/PDF justification
    pPr.append(jc)

    # 2. ENHANCED TEXT DISTRIBUTION for better spacing
    textAlignment = OxmlElement("w:textAlignment")
    textAlignment.set(qn("w:val"), "distribute")
    pPr.append(textAlignment)

    # 3. ENABLE AUTOMATIC HYPHENATION for better line breaks
    suppressAutoHyphens = OxmlElement("w:suppressAutoHyphens")
    suppressAutoHyphens.set(qn("w:val"), "0")  # Enable hyphenation
    pPr.append(suppressAutoHyphens)

    # 4. EXACT LINE SPACING (12pt = 240 twips) - MANDATORY EXACT RULE
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(spacing_before))
    spacing.set(qn("w:after"), str(spacing_after))
    spacing.set(qn("w:line"), str(line_spacing))
    spacing.set(qn("w:lineRule"), "exact")  # EXACT rule for perfect IEEE spacing
    pPr.append(spacing)

    # 5. ADVANCED JUSTIFICATION CONTROLS FOR PERFECT ALIGNMENT
    adjustRightInd = OxmlElement("w:adjustRightInd")
    adjustRightInd.set(qn("w:val"), "1")
    pPr.append(adjustRightInd)

    snapToGrid = OxmlElement("w:snapToGrid")
    snapToGrid.set(qn("w:val"), "0")
    pPr.append(snapToGrid)


def setup_two_column_layout(doc):
    """ENSURE COLUMNS APPLY AFTER ABSTRACT - Setup TWO-COLUMN LAYOUT with EXACT IEEE LaTeX specifications."""
    # Add section break BEFORE two columns (CRITICAL for proper column application)
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type = WD_SECTION.CONTINUOUS

    # Configure two-column layout via OpenXML with EXACT IEEE specifications
    sectPr = new_section._sectPr

    # Set margins - EXACT IEEE LaTeX: 0.75" all sides (1080 twips)
    pgMar = (
        sectPr.xpath("./w:pgMar")[0]
        if sectPr.xpath("./w:pgMar")
        else OxmlElement("w:pgMar")
    )
    pgMar.set(qn("w:left"), "1080")
    pgMar.set(qn("w:right"), "1080")
    pgMar.set(qn("w:top"), "1080")
    pgMar.set(qn("w:bottom"), "1080")
    if not sectPr.xpath("./w:pgMar"):
        sectPr.append(pgMar)

    # Remove existing cols element if present
    existing_cols = sectPr.xpath("./w:cols")
    for col in existing_cols:
        sectPr.remove(col)

    # Add new cols element with EXACT IEEE LaTeX specifications
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")  # 2 columns
    cols.set(qn("w:space"), "360")  # 0.25" gap (360 twips)
    cols.set(qn("w:equalWidth"), "1")  # Equal width columns

    # Add column definitions - EXACT 3.3125" per column (4770 twips)
    for i in range(2):
        col = OxmlElement("w:col")
        col.set(qn("w:w"), "4770")  # 3.3125" width
        cols.append(col)

    sectPr.append(cols)


def add_ieee_body_paragraph(doc, text):
    """Add a body paragraph with EXACT IEEE LaTeX PDF formatting via OpenXML."""
    para = doc.add_paragraph()
    run = para.add_run(sanitize_text(text))

    # Font: Times New Roman 10pt
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    # FIXED: Apply full IEEE justification using the dedicated function
    apply_ieee_latex_formatting(
        para, spacing_before=0, spacing_after=0, line_spacing=240
    )

    return para


def add_ieee_table(doc, table_data, section_idx, table_count):
    """COMPREHENSIVE TABLE VISIBILITY FIX - Ensures tables are 100% visible in Word documents."""
    try:
        table_type = table_data.get("tableType", table_data.get("type", "interactive"))
        print(f"🔧 Processing table type: {table_type}", file=sys.stderr)

        if table_type == "interactive":
            # STEP 1: COMPREHENSIVE DATA VALIDATION AND NORMALIZATION
            headers = table_data.get("headers", [])
            rows_data = table_data.get("tableData", []) or table_data.get("rows", [])
            
            print(f"📊 Table data validation:", file=sys.stderr)
            print(f"   Headers: {headers}", file=sys.stderr)
            print(f"   Rows: {len(rows_data)} rows", file=sys.stderr)
            print(f"   Available keys: {list(table_data.keys())}", file=sys.stderr)

            # CRITICAL: Ensure we have valid, visible table data
            if not headers or not rows_data:
                print(f"⚠️ Missing table data - creating visible placeholder", file=sys.stderr)
                headers = ["Parameter", "Value", "Description"]
                rows_data = [
                    ["Sample Parameter 1", "Sample Value 1", "Sample Description 1"],
                    ["Sample Parameter 2", "Sample Value 2", "Sample Description 2"],
                    ["Sample Parameter 3", "Sample Value 3", "Sample Description 3"]
                ]
            
            # Validate and clean data
            headers = [sanitize_text(str(h)) if h else f"Column {i+1}" for i, h in enumerate(headers)]
            cleaned_rows = []
            for row in rows_data:
                if isinstance(row, list):
                    cleaned_row = [sanitize_text(str(cell)) if cell else "" for cell in row]
                    # Ensure row has same number of columns as headers
                    while len(cleaned_row) < len(headers):
                        cleaned_row.append("")
                    cleaned_rows.append(cleaned_row[:len(headers)])  # Trim excess columns
                else:
                    # Handle non-list rows
                    cleaned_rows.append([sanitize_text(str(row))] + [""] * (len(headers) - 1))
            
            rows_data = cleaned_rows
            
            # STEP 2: PREPARE TABLE FOR 2-COLUMN LAYOUT COMPATIBILITY
            print("🔧 Preparing table for 2-column layout...", file=sys.stderr)
            
            # STEP 3: CREATE TABLE WITH MAXIMUM VISIBILITY SETTINGS
            num_cols = len(headers)
            num_rows = len(rows_data) + 1  # +1 for header row
            
            print(f"🔧 Creating table: {num_rows} rows × {num_cols} columns", file=sys.stderr)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            
            # STEP 4: APPLY MAXIMUM VISIBILITY TABLE FORMATTING
            table.style = None  # Remove any default style that might hide content
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.allow_autofit = False  # Disable autofit to prevent content hiding
            
            # STEP 5: SET COMPREHENSIVE TABLE PROPERTIES FOR MAXIMUM WORD COMPATIBILITY
            tbl = table._element
            tblPr = tbl.xpath('./w:tblPr')[0] if tbl.xpath('./w:tblPr') else OxmlElement('w:tblPr')
            
            # Clear any existing properties that might interfere with visibility
            for child in list(tblPr):
                tblPr.remove(child)
            
            # Set table width to fit within single column (2-column layout compatible)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '4770')  # Single column width in twips (3.3125" * 1440 twips/inch)
            tblW.set(qn('w:type'), 'dxa')  # Use exact measurements
            tblPr.append(tblW)
            
            # Center table alignment
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tblPr.append(jc)
            
            # Set table layout to fixed for consistent rendering
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # CRITICAL: Force table to never overlap with other content
            tblOverlap = OxmlElement('w:tblOverlap')
            tblOverlap.set(qn('w:val'), 'never')
            tblPr.append(tblOverlap)
            
            # Add table properties to table element
            if not tbl.xpath('./w:tblPr'):
                tbl.insert(0, tblPr)
            
            # STEP 6: SET COMPREHENSIVE CELL FORMATTING FOR MAXIMUM VISIBILITY
            print("🔧 Applying maximum visibility cell formatting...", file=sys.stderr)
            
            # Calculate column width for 2-column layout compatibility
            col_width_twips = 4770 // num_cols  # Distribute single column width equally
            min_width_twips = 720  # Minimum 0.5 inch per column for 2-column layout
            final_col_width = max(min_width_twips, col_width_twips)
            
            for row in table.rows:
                for cell in row.cells:
                    # Get cell element for direct XML manipulation
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Clear existing cell properties that might interfere
                    for child in list(tcPr):
                        tcPr.remove(child)
                    
                    # CRITICAL: Set visible borders for table structure
                    tcBorders = OxmlElement("w:tcBorders")
                    for border_name in ["top", "left", "bottom", "right"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "12")  # 1.5pt border for clear visibility
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "000000")  # Black border
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                    
                    # Set cell margins for proper content spacing
                    tcMar = OxmlElement("w:tcMar")
                    for margin_name in ["top", "left", "bottom", "right"]:
                        margin = OxmlElement(f"w:{margin_name}")
                        margin.set(qn("w:w"), "144")  # 0.1" margin for content visibility
                        margin.set(qn("w:type"), "dxa")
                        tcMar.append(margin)
                    tcPr.append(tcMar)
                    
                    # Set fixed cell width for consistent layout
                    tcW = OxmlElement("w:tcW")
                    tcW.set(qn("w:w"), str(final_col_width))
                    tcW.set(qn("w:type"), "dxa")
                    tcPr.append(tcW)
                    
                    # Force white background for maximum contrast
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFFF')  # White background
                    tcPr.append(shd)
                    
                    # Set vertical alignment to top for consistent appearance
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'top')
                    tcPr.append(vAlign)
            
            # STEP 7: POPULATE TABLE CONTENT WITH MAXIMUM VISIBILITY FORMATTING
            print("🔧 Populating table content with maximum visibility...", file=sys.stderr)
            
            # HEADER ROW - Bold, centered, maximum visibility
            header_row = table.rows[0]
            for col_idx, header in enumerate(headers):
                cell = header_row.cells[col_idx]
                
                # Clear existing content completely
                cell._element.clear_content()
                
                # Add new paragraph with maximum visibility
                para = cell.add_paragraph()
                run = para.add_run(str(header))
                
                # Maximum visibility header formatting
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)  # Standard size for maximum compatibility
                run.bold = True
                run.font.color.rgb = None  # Ensure black text
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Force paragraph visibility with proper spacing
                pPr = para._element.get_or_add_pPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:before"), "72")  # 0.05" before
                spacing.set(qn("w:after"), "72")   # 0.05" after
                spacing.set(qn("w:line"), "240")   # 12pt line height
                spacing.set(qn("w:lineRule"), "exact")
                pPr.append(spacing)
                
                print(f"✅ Added header: '{header}'", file=sys.stderr)

            # DATA ROWS - Regular, left-aligned, maximum visibility
            for row_idx, row_data in enumerate(rows_data):
                table_row = table.rows[row_idx + 1]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table_row.cells[col_idx]
                        
                        # Clear existing content completely
                        cell._element.clear_content()
                        
                        # Add new paragraph with maximum visibility
                        para = cell.add_paragraph()
                        run = para.add_run(str(cell_data))
                        
                        # Maximum visibility data formatting
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(9)  # Slightly smaller for data
                        run.bold = False
                        run.font.color.rgb = None  # Ensure black text
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Force paragraph visibility with proper spacing
                        pPr = para._element.get_or_add_pPr()
                        spacing = OxmlElement("w:spacing")
                        spacing.set(qn("w:before"), "36")  # Small spacing
                        spacing.set(qn("w:after"), "36")
                        spacing.set(qn("w:line"), "240")   # 12pt line height
                        spacing.set(qn("w:lineRule"), "exact")
                        pPr.append(spacing)
                        
                        print(f"✅ Added data cell [{row_idx}][{col_idx}]: '{cell_data}'", file=sys.stderr)

            # STEP 8: ADD SPACING AFTER TABLE (STAYS IN 2-COLUMN LAYOUT)
            print("🔧 Adding spacing after table within 2-column layout...", file=sys.stderr)
            
            # Add spacing after table
            spacing_after = doc.add_paragraph()
            spacing_after.paragraph_format.space_before = Pt(12)
            spacing_after.paragraph_format.space_after = Pt(12)
            
            print(f"🎉 Table {table_count} completed and fits in 2-column layout!", file=sys.stderr)
            return True
            if not tbl.xpath('./w:tblPr'):
                tbl.insert(0, tblPr)
            
            # CRITICAL: Ensure table has visible borders and proper cell formatting
            for row in table.rows:
                for cell in row.cells:
                    # Set cell borders to make table visible
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Remove any existing borders first
                    existing_borders = tcPr.xpath('./w:tcBorders')
                    for border in existing_borders:
                        tcPr.remove(border)
                    
                    # Add new visible borders
                    tcBorders = OxmlElement("w:tcBorders")
                    
                    # Add all borders (top, left, bottom, right) with thick black lines
                    for border_name in ["top", "left", "bottom", "right"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "18")  # 2.25pt thick border for maximum visibility
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "000000")  # Black border
                        tcBorders.append(border)
                    
                    tcPr.append(tcBorders)
                    
                    # CRITICAL: Ensure cell has proper margins and padding for content visibility
                    tcMar = OxmlElement("w:tcMar")
                    for margin_name in ["top", "left", "bottom", "right"]:
                        margin = OxmlElement(f"w:{margin_name}")
                        margin.set(qn("w:w"), "108")  # 0.075" margin for better content visibility
                        margin.set(qn("w:type"), "dxa")
                        tcMar.append(margin)
                    tcPr.append(tcMar)
                    
                    # Force cell to have minimum width for visibility
                    tcW = OxmlElement("w:tcW")
                    tcW.set(qn("w:w"), "2160")  # Use fixed width for consistency
                    tcW.set(qn("w:type"), "dxa")
                    tcPr.append(tcW)

            # Set column widths for 2-column layout compatibility (integer values required)
            available_width_twips = 4770  # Single column width in twips (3.3125")
            col_width_twips = available_width_twips // num_cols  # Integer division
            min_width_twips = 720  # 0.5 inch minimum in twips for 2-column layout
            
            final_col_width = max(min_width_twips, col_width_twips)
            
            for col in table.columns:
                col.width = final_col_width
                
            print(f"Set column width to {final_col_width} twips for {num_cols} columns (2-column layout)", file=sys.stderr)

            # HEADER ROW - Bold, centered, 9pt Times New Roman with MAXIMUM VISIBILITY
            header_row = table.rows[0]
            for col_idx, header in enumerate(headers):
                cell = header_row.cells[col_idx]
                
                # Clear existing content and add new content
                cell._element.clear_content()
                para = cell.add_paragraph()
                run = para.add_run(sanitize_text(str(header)))
                
                # MAXIMUM VISIBILITY header formatting
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)  # Slightly larger for visibility
                run.bold = True  # BOLD headers
                run.font.color.rgb = None  # Ensure black text
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Force paragraph to be visible
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                
                print(f"Added header: {header}", file=sys.stderr)

            # DATA ROWS - Regular, left-aligned, 9pt Times New Roman with MAXIMUM VISIBILITY
            for row_idx, row_data in enumerate(rows_data):
                table_row = table.rows[row_idx + 1]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table_row.cells[col_idx]
                        
                        # Clear existing content and add new content
                        cell._element.clear_content()
                        para = cell.add_paragraph()
                        run = para.add_run(sanitize_text(str(cell_data)))
                        
                        # MAXIMUM VISIBILITY data formatting
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)  # Slightly larger for visibility
                        run.bold = False  # Regular weight for data
                        run.font.color.rgb = None  # Ensure black text
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Force paragraph to be visible
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        
                        # CRITICAL: Ensure cell has proper background
                        tcPr = cell._element.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'FFFFFF')  # White background
                        tcPr.append(shd)
                        
                        print(f"Added data cell [{row_idx}][{col_idx}]: {cell_data}", file=sys.stderr)

            # TABLE SPACING - 6pt before/after (EXACT IEEE specification)
            # Add spacing paragraph after table (table is already added to document)
            spacing_after = doc.add_paragraph()
            spacing_after.paragraph_format.space_before = Pt(6)
            spacing_after.paragraph_format.space_after = Pt(6)

        elif table_type == "image":
            # Handle image tables within 2-column layout
            if table_data.get("data"):
                try:
                    print("🔧 Processing image table for 2-column layout...", file=sys.stderr)
                    
                    image_data = table_data["data"]
                    if "," in image_data:
                        image_data = image_data.split(",")[1]

                    image_bytes = base64.b64decode(image_data)
                    image_stream = BytesIO(image_bytes)

                    # Add spacing before image
                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_before = Pt(12)
                    spacing_para.paragraph_format.space_after = Pt(6)

                    # Create image paragraph with center alignment
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.paragraph_format.space_before = Pt(6)
                    img_para.paragraph_format.space_after = Pt(6)

                    # Add image to paragraph with maximum visibility settings
                    run = img_para.add_run()

                    # Size for 2-column layout compatibility - CONSERVATIVE SIZING TO PREVENT OVERLAP
                    size = table_data.get("size", "medium")
                    size_mapping = {
                        "small": Inches(1.0),   # Very conservative for 2-column
                        "medium": Inches(1.4),  # Conservative for 2-column
                        "large": Inches(1.8),   # Max conservative for 2-column
                    }
                    width = size_mapping.get(size, Inches(1.4))
                    
                    print(f"📏 Image table size '{size}' mapped to width: {width}", file=sys.stderr)

                    # Add image with comprehensive visibility settings
                    picture = run.add_picture(image_stream, width=width)
                    
                    # CRITICAL: Set image positioning properties for maximum visibility
                    inline = picture._inline
                    extent = inline.extent
                    
                    # Force image to be visible and properly positioned
                    docPr = inline.docPr
                    docPr.set('name', f'Image_{table_count}')
                    docPr.set('descr', 'Table Image')
                    
                    # Set graphic properties for maximum visibility
                    graphic = inline.graphic
                    graphicData = graphic.graphicData
                    pic = graphicData.pic
                    
                    # Ensure image has proper display properties
                    spPr = pic.spPr
                    if spPr is not None:
                        # Add transform properties for proper positioning
                        xfrm = spPr.xfrm
                        if xfrm is not None:
                            # Ensure image is not clipped
                            off = xfrm.off
                            if off is not None:
                                off.set('x', '0')
                                off.set('y', '0')
                    
                    # Scale down if too tall to prevent page overflow
                    if picture.height > IEEE_CONFIG["max_figure_height"]:
                        scale_factor = IEEE_CONFIG["max_figure_height"] / picture.height
                        run.clear()
                        image_stream.seek(0)
                        picture = run.add_picture(
                            image_stream,
                            width=width * scale_factor,
                            height=IEEE_CONFIG["max_figure_height"],
                        )
                        
                        # Reapply visibility settings after scaling
                        inline = picture._inline
                        docPr = inline.docPr
                        docPr.set('name', f'Image_{table_count}_scaled')
                        docPr.set('descr', 'Scaled Table Image')

                    # Add spacing after image
                    spacing_para_after = doc.add_paragraph()
                    spacing_para_after.paragraph_format.space_before = Pt(6)
                    spacing_para_after.paragraph_format.space_after = Pt(12)
                    
                    print("✅ Image table processed for 2-column layout", file=sys.stderr)

                except Exception as e:
                    print(f"Error processing table image: {e}", file=sys.stderr)
                    return False

        elif table_type == "latex":
            # Handle LaTeX tables within 2-column layout
            latex_code = table_data.get("latexCode", "")
            if latex_code:
                print("🔧 Processing LaTeX table for 2-column layout...", file=sys.stderr)
                
                # Add LaTeX code with proper formatting (stays in 2-column layout)
                para = doc.add_paragraph()
                run = para.add_run(f"LaTeX Table Code:\n{sanitize_text(latex_code)}")
                run.font.name = "Courier New"
                run.font.size = Pt(8)  # Smaller font for 2-column layout
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                
                print("✅ LaTeX table processed for 2-column layout", file=sys.stderr)

        return True

    except Exception as e:
        print(f"Error adding table: {e}", file=sys.stderr)
        import traceback

        traceback.print_exc(file=sys.stderr)


def add_justified_paragraph(
    doc,
    text,
    style_name="Normal",
    indent_left=None,
    indent_right=None,
    space_before=None,
    space_after=None,
):
    """Add a paragraph with professional justification settings for research paper quality."""
    # Use the new IEEE body paragraph function for exact formatting
    return add_ieee_body_paragraph(doc, text)


def add_section(doc, section_data, section_idx, is_first_section=False):
    """Add a section with content blocks, subsections, and figures - EXACT same as test.py."""
    if section_data.get("title"):
        # Create section heading with exact IEEE LaTeX formatting
        para = doc.add_paragraph()
        run = para.add_run(
            f"{section_idx}. {sanitize_text(section_data['title']).upper()}"
        )

        # Font: Times New Roman 10pt bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True

        # Apply exact OpenXML formatting for section headings
        pPr = para._element.get_or_add_pPr()

        # Clear existing formatting
        for elem in pPr.xpath("./w:spacing | ./w:jc"):
            pPr.remove(elem)

        # Center alignment (IEEE standard for section headings)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)

        # Section heading spacing: 12pt before, 0pt after
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "240")  # 12pt before
        spacing.set(qn("w:after"), "0")  # 0pt after
        spacing.set(qn("w:line"), "240")  # 12pt line spacing
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

    # Process content blocks (text and images in order) - Support BOTH naming conventions
    content_blocks = section_data.get("contentBlocks", []) or section_data.get(
        "content_blocks", []
    )

    # Track table count for numbering
    table_count = 0

    for block_idx, block in enumerate(content_blocks):
        if block.get("type") == "text" and block.get("content"):
            space_before = (
                IEEE_CONFIG["line_spacing"]
                if is_first_section and block_idx == 0
                else Pt(3)
            )
            add_formatted_paragraph(
                doc,
                block["content"],
                indent_left=IEEE_CONFIG["column_indent"],
                indent_right=IEEE_CONFIG["column_indent"],
                space_before=space_before,
                space_after=Pt(12),
            )

        elif block.get("type") == "table":
            # Handle table blocks with comprehensive visibility fixes
            table_count += 1
            
            print(f"🔧 Processing table block in section {section_idx}, table {table_count}", file=sys.stderr)
            print(f"📊 Table block data: {block}", file=sys.stderr)

            # COMPREHENSIVE CAPTION HANDLING - Support all possible caption fields
            caption_text = (
                block.get("caption", "").strip() or 
                block.get("tableName", "").strip() or 
                block.get("name", "").strip() or 
                f"Data Table {table_count}"
            )

            print(f"📝 Final table caption: {caption_text}", file=sys.stderr)

            # Add table caption BEFORE table with proper formatting
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(f"TABLE {section_idx}.{table_count}: {sanitize_text(caption_text).upper()}")
            caption_run.font.name = "Times New Roman"
            caption_run.font.size = Pt(9)
            caption_run.bold = True
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.paragraph_format.space_before = Pt(12)
            caption_para.paragraph_format.space_after = Pt(6)

            # Add the table content with comprehensive visibility fixes
            success = add_ieee_table(doc, block, section_idx, table_count)
            
            if success:
                print(f"✅ Successfully completed table {table_count} in section {section_idx}", file=sys.stderr)
            else:
                print(f"❌ Failed to create table {table_count} in section {section_idx}", file=sys.stderr)

        elif block.get("type") == "text" and block.get("data") and block.get("caption"):
            # Handle text blocks with attached images (React frontend pattern)
                # Handle image attached to text block
                import base64

                size = block.get("size", "medium")
                # Map frontend size names to backend size names
                size_mapping = {
                    "very-small": "Very Small",
                    "small": "Small",
                    "medium": "Medium",
                    "large": "Large",
                }
                mapped_size = size_mapping.get(size, "Medium")
                width = IEEE_CONFIG["figure_sizes"].get(
                    mapped_size, IEEE_CONFIG["figure_sizes"]["Medium"]
                )

                # Decode base64 image data
                try:
                    image_data = block["data"]

                    # Handle base64 data - remove prefix if present
                    if "," in image_data:
                        image_data = image_data.split(",")[1]

                    # Decode base64 image data
                    try:
                        image_bytes = base64.b64decode(image_data)
                    except Exception as e:
                        print(
                            f"ERROR: Failed to decode image data in text block: {str(e)}",
                            file=sys.stderr,
                        )
                        continue

                    # Create image stream
                    image_stream = BytesIO(image_bytes)

                    # Simple image layout without complex section breaks
                    # Add simple spacing before image
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
                    
                    # Create simple image paragraph
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add the image with minimal formatting
                    run = para.add_run()
                    
                    try:
                        # Add image with proper sizing
                        picture = run.add_picture(image_stream, width=width)
                        
                        # Scale down if too tall to prevent page overflow
                        if picture.height > IEEE_CONFIG["max_figure_height"]:
                            scale_factor = IEEE_CONFIG["max_figure_height"] / picture.height
                            run.clear()
                            image_stream.seek(0)  # Reset stream position
                            picture = run.add_picture(
                                image_stream,
                                width=width * scale_factor,
                                height=IEEE_CONFIG["max_figure_height"],
                            )
                        
                        print(f"Added image with simplified layout, width: {width}, height: {picture.height}", file=sys.stderr)
                        
                    except Exception as img_error:
                        print(f"Error adding image: {img_error}", file=sys.stderr)
                        # Fallback: add text placeholder
                        run.add_text(f"[Image: {block.get('caption', 'Figure')}]")

                    # Generate figure number and add simple caption
                    img_count = sum(
                        1
                        for b in content_blocks[: block_idx + 1]
                        if b.get("type") == "image"
                    )
                    
                    # Add figure caption with simple formatting
                    caption_para = doc.add_paragraph()
                    caption_run = caption_para.add_run(f"FIG. {section_idx}.{img_count}: {sanitize_text(block['caption']).upper()}")
                    caption_run.font.name = "Times New Roman"
                    caption_run.font.size = Pt(9)
                    caption_run.bold = True
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add simple spacing after caption
                    doc.add_paragraph().paragraph_format.space_after = Pt(12)
                except Exception as e:
                    print(f"Error processing image in text block: {e}", file=sys.stderr)

        elif (
            block.get("type") == "image" and block.get("data") and block.get("caption")
        ):
            # Handle image blocks with proper ordering and visibility
            img_count = sum(
                1 for b in content_blocks[: block_idx + 1] if b.get("type") == "image"
            )

            print(f"Processing image block in section {section_idx}, image {img_count}", file=sys.stderr)

            # Get image size and caption - CONSERVATIVE SIZING TO PREVENT OVERLAP
            size = block.get("size", "medium")
            size_mapping = {
                "very-small": Inches(1.0),  # Very conservative for 2-column
                "small": Inches(1.4),       # Conservative for 2-column
                "medium": Inches(1.8),      # Conservative for 2-column
                "large": Inches(2.2),       # Max conservative for 2-column
            }
            width = size_mapping.get(size, Inches(1.8))
            
            print(f"📏 Image size '{size}' mapped to width: {width}", file=sys.stderr)
            caption_text = sanitize_text(block['caption'])
            figure_number = f"FIG. {section_idx}.{img_count}"

            # SIMPLE & EFFECTIVE: Create properly sized inline image for 2-column layout
            try:
                import base64
                # Decode base64 image data
                image_data = block["data"]
                if "," in image_data:
                    image_data = image_data.split(",")[1]
                image_bytes = base64.b64decode(image_data)
                image_stream = BytesIO(image_bytes)

                print(f"🔧 Processing image {figure_number} for 2-column layout compatibility...", file=sys.stderr)

                # ULTIMATE FIX: Force column break before image to ensure complete separation
                from docx.enum.text import WD_BREAK
                
                # Add paragraph with column break to force image to new position
                break_para = doc.add_paragraph()
                break_run = break_para.add_run()
                break_run.add_break(WD_BREAK.COLUMN)
                break_para.paragraph_format.space_before = Pt(0)
                break_para.paragraph_format.space_after = Pt(6)

                # FORCE LINE BREAK BEFORE IMAGE - Prevents inline text issues
                line_break_before = doc.add_paragraph()
                line_break_before.paragraph_format.space_before = Pt(6)
                line_break_before.paragraph_format.space_after = Pt(3)
                line_break_before.paragraph_format.keep_with_next = False
                line_break_before.paragraph_format.page_break_before = False

                # Create image paragraph with FORCED isolation from text
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.paragraph_format.space_before = Pt(6)
                img_para.paragraph_format.space_after = Pt(6)
                img_para.paragraph_format.keep_with_next = True  # Keep image with caption
                img_para.paragraph_format.keep_together = True  # Prevent splitting
                
                # Set paragraph properties to prevent clipping in 2-column layout
                pPr = img_para._element.get_or_add_pPr()
                
                # Clear any existing formatting that might interfere
                for elem in pPr.xpath("./w:spacing | ./w:ind | ./w:jc | ./w:keepNext | ./w:keepLines"):
                    pPr.remove(elem)
                
                # Center alignment for images
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                pPr.append(jc)
                
                # CRITICAL: Force paragraph to be on its own line
                keepNext = OxmlElement("w:keepNext")
                keepNext.set(qn("w:val"), "1")
                pPr.append(keepNext)
                
                keepLines = OxmlElement("w:keepLines")
                keepLines.set(qn("w:val"), "1")
                pPr.append(keepLines)
                
                # Proper spacing for image visibility with generous margins
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:before"), "240")  # 12pt before (increased)
                spacing.set(qn("w:after"), "240")   # 12pt after (increased)
                spacing.set(qn("w:line"), "360")    # 18pt line height (increased)
                spacing.set(qn("w:lineRule"), "exact")
                pPr.append(spacing)
                
                # Ensure no indentation that might cause clipping
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "0")
                ind.set(qn("w:right"), "0")
                pPr.append(ind)
                
                # Add image with proper sizing for 2-column layout
                run = img_para.add_run()
                
                # Ensure image fits comfortably within column width with generous margins
                max_col_width = Inches(2.2)  # Much smaller to prevent overlap issues
                if width > max_col_width:
                    width = max_col_width
                
                print(f"📏 Image width set to: {width} (max allowed: {max_col_width})", file=sys.stderr)
                
                picture = run.add_picture(image_stream, width=width)
                
                # Set comprehensive image properties for maximum visibility
                inline = picture._inline
                docPr = inline.docPr
                docPr.set('name', f'Figure_{section_idx}_{img_count}')
                docPr.set('descr', f'Section Figure: {caption_text}')
                
                # Ensure image has proper positioning within the paragraph
                graphic = inline.graphic
                graphicData = graphic.graphicData
                pic = graphicData.pic
                
                # Set image properties for maximum visibility
                spPr = pic.spPr
                if spPr is not None:
                    xfrm = spPr.xfrm
                    if xfrm is not None:
                        # Ensure image is properly positioned
                        off = xfrm.off
                        if off is not None:
                            off.set('x', '0')
                            off.set('y', '0')
                
                # Scale down if too tall
                if picture.height > IEEE_CONFIG["max_figure_height"]:
                    scale_factor = IEEE_CONFIG["max_figure_height"] / picture.height
                    run.clear()
                    image_stream.seek(0)
                    picture = run.add_picture(
                        image_stream,
                        width=width * scale_factor,
                        height=IEEE_CONFIG["max_figure_height"],
                    )
                    
                    # Reapply properties after scaling
                    inline = picture._inline
                    docPr = inline.docPr
                    docPr.set('name', f'Figure_{section_idx}_{img_count}_scaled')
                    docPr.set('descr', f'Section Figure: {caption_text} (scaled)')

                # Add figure caption in separate paragraph
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"{figure_number}: {caption_text.upper()}")
                caption_run.font.name = "Times New Roman"
                caption_run.font.size = Pt(9)
                caption_run.bold = True
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_before = Pt(3)
                caption_para.paragraph_format.space_after = Pt(12)

                # FORCE COMPLETE SEPARATION - Prevents any inline text issues
                spacing_after = doc.add_paragraph()
                spacing_after.paragraph_format.space_before = Pt(6)
                spacing_after.paragraph_format.space_after = Pt(24)  # Large spacing to prevent overlap
                spacing_after.paragraph_format.keep_with_next = False
                spacing_after.paragraph_format.page_break_before = False
                
                # Add mandatory paragraph break to ensure complete separation
                separator_para = doc.add_paragraph()
                separator_para.paragraph_format.space_before = Pt(0)
                separator_para.paragraph_format.space_after = Pt(18)
                separator_para.paragraph_format.keep_with_next = False
                
                # Add final spacing paragraph to guarantee text separation
                final_spacing = doc.add_paragraph()
                final_spacing.paragraph_format.space_before = Pt(0)
                final_spacing.paragraph_format.space_after = Pt(12)
                final_spacing.paragraph_format.keep_with_next = False

                print(f"✅ Successfully added image {figure_number} optimized for 2-column layout", file=sys.stderr)

            except Exception as e:
                print(f"❌ Error adding image {figure_number}: {e}", file=sys.stderr)
                # Add simple placeholder if image fails
                para = doc.add_paragraph(f"[Image: {caption_text}]")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif block.get("type") == "equation" and block.get("content"):
            # Handle equation blocks
            equation_content = sanitize_text(block.get("content", ""))
            equation_number = block.get("equationNumber", "")
            
            # Add equation paragraph with center alignment
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            
            # Add equation content with number if provided
            if equation_number:
                run = para.add_run(f"({equation_number}) {equation_content}")
            else:
                run = para.add_run(equation_content)
            
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True  # IEEE equations are typically italic

        elif block.get("type") == "subsection":
            # Handle subsection blocks
            subsection_title = sanitize_text(block.get("title", ""))
            subsection_content = sanitize_text(block.get("content", ""))
            
            if subsection_title:
                # Add subsection heading
                para = doc.add_heading(subsection_title, level=3)
                para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.keep_with_next = True
            
            if subsection_content:
                # Add subsection content
                add_formatted_paragraph(
                    doc,
                    subsection_content,
                    indent_left=IEEE_CONFIG["column_indent"],
                    indent_right=IEEE_CONFIG["column_indent"],
                    space_before=Pt(3),
                    space_after=Pt(12),
                )

    # Legacy support for old content field - EXACT same as test.py
    if not content_blocks and section_data.get("content"):
        space_before = IEEE_CONFIG["line_spacing"] if is_first_section else Pt(3)
        add_justified_paragraph(
            doc,
            section_data["content"],
            indent_left=IEEE_CONFIG["column_indent"],
            indent_right=IEEE_CONFIG["column_indent"],
            space_before=space_before,
            space_after=Pt(12),
        )

    # Add subsections with multi-level support
    def add_subsection_recursive(subsections, section_idx, parent_numbering=""):
        """Recursively add subsections with proper hierarchical numbering."""
        # Group subsections by level and parent
        level_1_subsections = [
            s for s in subsections if s.get("level", 1) == 1 and not s.get("parentId")
        ]

        for sub_idx, subsection in enumerate(level_1_subsections, 1):
            if subsection.get("title"):
                subsection_number = f"{section_idx}.{sub_idx}"
                para = doc.add_heading(
                    f"{subsection_number} {sanitize_text(subsection['title'])}", level=2
                )
                para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = IEEE_CONFIG["line_spacing"]
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = False
                para.paragraph_format.widow_control = False

            if subsection.get("content"):
                add_justified_paragraph(
                    doc,
                    sanitize_text(subsection["content"]),
                    indent_left=IEEE_CONFIG["column_indent"],
                    indent_right=IEEE_CONFIG["column_indent"],
                    space_before=Pt(1),
                    space_after=Pt(12),
                )

            # Handle nested subsections (level 2 and beyond)
            add_nested_subsection(
                subsections, subsection["id"], f"{section_idx}.{sub_idx}", 2
            )

    def add_nested_subsection(all_subsections, parent_id, parent_number, level):
        """Add nested subsections recursively."""
        child_subsections = [
            s
            for s in all_subsections
            if s.get("parentId") == parent_id and s.get("level", 1) == level
        ]

        for child_idx, child_sub in enumerate(child_subsections, 1):
            # Always define child_number, regardless of whether title exists
            child_number = f"{parent_number}.{child_idx}"

            if child_sub.get("title"):
                # Use different heading levels for deeper nesting, but cap at level 6
                heading_level = min(level + 1, 6)
                para = doc.add_heading(
                    f"{child_number} {sanitize_text(child_sub['title'])}",
                    level=heading_level,
                )
                para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = False
                para.paragraph_format.widow_control = False

            if child_sub.get("content"):
                add_justified_paragraph(
                    doc,
                    sanitize_text(child_sub["content"]),
                    indent_left=IEEE_CONFIG["column_indent"]
                    + Inches(0.1 * (level - 1)),  # Progressive indentation
                    indent_right=IEEE_CONFIG["column_indent"],
                    space_before=Pt(1),
                    space_after=Pt(12),
                )

            # Process content blocks if they exist
            if child_sub.get("contentBlocks"):
                for block in child_sub["contentBlocks"]:
                    if block.get("type") == "text" and block.get("content"):
                        add_formatted_paragraph(
                            doc,
                            block["content"],
                            indent_left=IEEE_CONFIG["column_indent"]
                            + Inches(0.1 * (level - 1)),
                            indent_right=IEEE_CONFIG["column_indent"],
                            space_before=Pt(1),
                            space_after=Pt(12),
                        )

            # Recursively handle even deeper nesting
            if level < 5:  # Limit depth to prevent excessive nesting
                add_nested_subsection(
                    all_subsections, child_sub["id"], child_number, level + 1
                )

    # Call the recursive function to add all subsections
    add_subsection_recursive(section_data.get("subsections", []), section_idx)


class HTMLToWordParser(HTMLParser):
    """Parse HTML content and apply formatting to Word document."""

    def __init__(self, paragraph):
        super().__init__()
        self.paragraph = paragraph
        self.format_stack = []
        self.text_buffer = ""

    def handle_starttag(self, tag, attrs):
        # Flush any buffered text before starting new formatting
        self._flush_text()

        if tag.lower() in ["b", "strong"]:
            self.format_stack.append("bold")
        elif tag.lower() in ["i", "em"]:
            self.format_stack.append("italic")
        elif tag.lower() == "u":
            self.format_stack.append("underline")

    def handle_endtag(self, tag):
        # Flush any buffered text before ending formatting
        self._flush_text()

        if tag.lower() in ["b", "strong"] and "bold" in self.format_stack:
            self.format_stack.remove("bold")
        elif tag.lower() in ["i", "em"] and "italic" in self.format_stack:
            self.format_stack.remove("italic")
        elif tag.lower() == "u" and "underline" in self.format_stack:
            self.format_stack.remove("underline")

    def handle_data(self, data):
        # Buffer the text data with sanitization
        self.text_buffer += sanitize_text(data)

    def _flush_text(self):
        """Create a run with accumulated text and current formatting."""
        if self.text_buffer:
            run = self.paragraph.add_run(self.text_buffer)
            run.font.name = IEEE_CONFIG["font_name"]
            run.font.size = IEEE_CONFIG["font_size_body"]

            # Apply current formatting
            if "bold" in self.format_stack:
                run.bold = True
            if "italic" in self.format_stack:
                run.italic = True
            if "underline" in self.format_stack:
                run.underline = True

            self.text_buffer = ""

    def close(self):
        """Ensure any remaining text is flushed when parsing is complete."""
        self._flush_text()
        super().close()


def apply_equal_justification(para):
    """Apply comprehensive equal justification controls for perfect equal line lengths like research papers."""
    # Get paragraph element for XML manipulation
    para_element = para._element
    pPr = para_element.get_or_add_pPr()

    # Primary justification setting - use 'both' for consistent Word/PDF output
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")  # Use 'both' for consistent Word/PDF justification
    pPr.append(jc)

    # Proper line spacing for 10pt text (IEEE standard)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:line"), "276")  # 13.8pt line spacing in twips (1.38 * 10pt * 20)
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)

    # Text alignment baseline for consistent line positioning
    textAlignment = OxmlElement("w:textAlignment")
    textAlignment.set(qn("w:val"), "baseline")
    pPr.append(textAlignment)

    # Force equal line lengths by enabling right margin adjustment
    adjust_right_ind = OxmlElement("w:adjustRightInd")
    adjust_right_ind.set(
        qn("w:val"), "1"
    )  # Allow right margin adjustment for equal line lengths
    pPr.append(adjust_right_ind)

    # Enable text compression for better line fitting
    compress_punctuation = OxmlElement("w:compressPunctuation")
    compress_punctuation.set(qn("w:val"), "1")
    pPr.append(compress_punctuation)

    # Control automatic spacing for better justification
    auto_space_de = OxmlElement("w:autoSpaceDE")
    auto_space_de.set(
        qn("w:val"), "1"
    )  # Enable auto spacing between Asian and Latin text
    pPr.append(auto_space_de)

    auto_space_dn = OxmlElement("w:autoSpaceDN")
    auto_space_dn.set(
        qn("w:val"), "1"
    )  # Enable auto spacing between Asian text and numbers
    pPr.append(auto_space_dn)

    # Word wrap settings for better line breaks
    word_wrap = OxmlElement("w:wordWrap")
    word_wrap.set(qn("w:val"), "1")
    pPr.append(word_wrap)

    # Enable text distribution for equal line lengths
    text_direction = OxmlElement("w:textDirection")
    text_direction.set(qn("w:val"), "lrTb")  # Left-to-right, top-to-bottom
    pPr.append(text_direction)

    # Force equal line distribution
    snap_to_grid = OxmlElement("w:snapToGrid")
    snap_to_grid.set(qn("w:val"), "0")  # Disable grid snapping for better justification
    pPr.append(snap_to_grid)

    # Character-level spacing controls for runs - allow flexibility for equal lines
    for run in para.runs:
        run_element = run._element
        rPr = run_element.get_or_add_rPr()

        # Allow character spacing adjustment for equal line lengths
        char_spacing = OxmlElement("w:spacing")
        char_spacing.set(qn("w:val"), "0")  # Start with no compression
        rPr.append(char_spacing)

        # Enable kerning for professional typography
        kern = OxmlElement("w:kern")
        kern.set(qn("w:val"), "20")  # 1pt kerning threshold
        rPr.append(kern)

        # Allow position adjustment for better line fitting
        position = OxmlElement("w:position")
        position.set(qn("w:val"), "0")
        rPr.append(position)

        # Enable text scaling for equal line lengths
        w_element = OxmlElement("w:w")
        w_element.set(
            qn("w:val"), "100"
        )  # 100% width scaling (can be adjusted by Word)
        rPr.append(w_element)

    return para


def apply_perfect_research_justification(para):
    """Apply PERFECT research paper quality justification - 200% perfect like top academic journals."""

    # Get paragraph element for XML manipulation
    para_element = para._element
    pPr = para_element.get_or_add_pPr()

    # AGGRESSIVE justification setting - use 'both' for consistent Word/PDF output
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")  # Use 'both' for consistent Word/PDF justification
    pPr.append(jc)

    # PERFECT line spacing for 10pt text (IEEE standard) - EXACT control
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:line"), "276")  # 13.8pt line spacing in twips (1.38 * 10pt * 20)
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)

    # BASELINE alignment for research paper consistency
    textAlignment = OxmlElement("w:textAlignment")
    textAlignment.set(qn("w:val"), "baseline")
    pPr.append(textAlignment)

    # FORCE perfect line lengths by AGGRESSIVE right margin adjustment
    adjust_right_ind = OxmlElement("w:adjustRightInd")
    adjust_right_ind.set(
        qn("w:val"), "1"
    )  # MAXIMUM right margin adjustment for perfect equal line lengths
    pPr.append(adjust_right_ind)

    # MAXIMUM text compression for perfect line fitting
    compress_punctuation = OxmlElement("w:compressPunctuation")
    compress_punctuation.set(qn("w:val"), "1")
    pPr.append(compress_punctuation)

    # PERFECT spacing controls for research paper quality
    auto_space_de = OxmlElement("w:autoSpaceDE")
    auto_space_de.set(qn("w:val"), "1")  # Perfect Asian-Latin spacing
    pPr.append(auto_space_de)

    auto_space_dn = OxmlElement("w:autoSpaceDN")
    auto_space_dn.set(qn("w:val"), "1")  # Perfect Asian-numeric spacing
    pPr.append(auto_space_dn)

    # PROFESSIONAL word wrapping for academic quality
    word_wrap = OxmlElement("w:wordWrap")
    word_wrap.set(qn("w:val"), "1")
    pPr.append(word_wrap)

    # PERFECT text direction for research papers
    text_direction = OxmlElement("w:textDirection")
    text_direction.set(qn("w:val"), "lrTb")  # Left-to-right, top-to-bottom
    pPr.append(text_direction)

    # DISABLE grid snapping for MAXIMUM justification control
    snap_to_grid = OxmlElement("w:snapToGrid")
    snap_to_grid.set(
        qn("w:val"), "0"
    )  # No grid interference with perfect justification
    pPr.append(snap_to_grid)

    # RESEARCH PAPER specific advanced controls
    # Force consistent line heights for academic quality
    line_rule = OxmlElement("w:lineRule")
    line_rule.set(qn("w:val"), "exact")
    pPr.append(line_rule)

    # Prevent widow/orphan that breaks perfect justification
    widow_control = OxmlElement("w:widowControl")
    widow_control.set(qn("w:val"), "0")
    pPr.append(widow_control)

    # ACADEMIC JOURNAL quality spacing - no auto spacing
    space_before_auto = OxmlElement("w:spaceBeforeAuto")
    space_before_auto.set(qn("w:val"), "0")
    pPr.append(space_before_auto)

    space_after_auto = OxmlElement("w:spaceAfterAuto")
    space_after_auto.set(qn("w:val"), "0")
    pPr.append(space_after_auto)

    # AGGRESSIVE character-level controls for PERFECT distribution
    for run in para.runs:
        run_element = run._element
        rPr = run_element.get_or_add_rPr()

        # AGGRESSIVE character spacing for perfect equal line lengths
        char_spacing = OxmlElement("w:spacing")
        char_spacing.set(
            qn("w:val"), "-20"
        )  # STRONG compression for perfect distribution
        rPr.append(char_spacing)

        # MAXIMUM kerning for professional typography
        kern = OxmlElement("w:kern")
        kern.set(
            qn("w:val"), "14"
        )  # Aggressive kerning for tight, professional spacing
        rPr.append(kern)

        # PRECISE position control for perfect alignment
        position = OxmlElement("w:position")
        position.set(qn("w:val"), "0")
        rPr.append(position)

        # AGGRESSIVE text scaling for perfect line fitting
        w_element = OxmlElement("w:w")
        w_element.set(
            qn("w:val"), "90"
        )  # 90% width scaling for tighter, more professional fit
        rPr.append(w_element)

        # RESEARCH PAPER quality font controls
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(
            qn("w:val"), str(int(IEEE_CONFIG["font_size_body"].pt * 2))
        )  # Ensure consistent size
        rPr.append(sz_cs)

        # DISABLE automatic expansion that breaks perfect justification
        no_proof = OxmlElement("w:noProof")
        no_proof.set(qn("w:val"), "1")
        rPr.append(no_proof)

    return para


def add_formatted_paragraph(doc, html_content, **kwargs):
    """FIXED: Add a paragraph with HTML formatting support and EXACT IEEE LaTeX formatting."""
    para = add_ieee_body_paragraph(doc, html_content or "")

    # FIXED: Ensure spacing parameters are applied if provided
    if "space_before" in kwargs and kwargs["space_before"] is not None:
        para.paragraph_format.space_before = kwargs["space_before"]
    if "space_after" in kwargs and kwargs["space_after"] is not None:
        para.paragraph_format.space_after = kwargs["space_after"]

    return para


def add_references(doc, references):
    """Add references section with proper alignment (hanging indent) and perfect justification."""
    if references:
        para = doc.add_heading("REFERENCES", level=1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center references heading
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False

        for idx, ref in enumerate(references, 1):
            # Handle both string references and object references
            if isinstance(ref, str):
                ref_text = sanitize_text(ref)
            elif isinstance(ref, dict) and ref.get("text"):
                ref_text = sanitize_text(ref["text"])
            else:
                continue  # Skip invalid references

            # Create reference paragraph with hanging indent
            para = doc.add_paragraph(f"[{idx}] {ref_text}")

            # Apply IEEE reference formatting with hanging indent
            pPr = para._element.get_or_add_pPr()

            # Hanging indent: 0.25" (360 twips)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:hanging"), "360")  # 0.25" hanging indent
            pPr.append(ind)

            # Reference spacing: 3pt before, 12pt after, 10pt line spacing (IEEE standard)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "60")  # 3pt before
            spacing.set(qn("w:after"), "240")  # 12pt after
            spacing.set(
                qn("w:line"), "200"
            )  # 10pt line spacing for references (IEEE standard)
            spacing.set(qn("w:lineRule"), "exact")
            pPr.append(spacing)

            # Set font: Times New Roman 9pt
            if para.runs:
                para.runs[0].font.name = "Times New Roman"
                para.runs[0].font.size = Pt(9)  # 9pt for references

                # Apply perfect justification with equal line lengths
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                apply_equal_justification(para)


def enable_auto_hyphenation(doc):
    """Enable professional hyphenation to improve justification quality."""
    section = doc.sections[-1]
    sectPr = section._sectPr

    # Enable automatic hyphenation for better justification
    auto_hyphenation = OxmlElement("w:autoHyphenation")
    auto_hyphenation.set(qn("w:val"), "1")
    sectPr.append(auto_hyphenation)

    # Do NOT hyphenate capitalized words (proper nouns, acronyms)
    do_not_hyphenate_caps = OxmlElement("w:doNotHyphenateCaps")
    do_not_hyphenate_caps.set(qn("w:val"), "1")
    sectPr.append(do_not_hyphenate_caps)

    # Set optimal hyphenation zone for research papers (0.25 inch)
    hyphenation_zone = OxmlElement("w:hyphenationZone")
    hyphenation_zone.set(qn("w:val"), "360")  # 0.25 inch in twips
    sectPr.append(hyphenation_zone)

    # Limit consecutive hyphens to maintain readability
    consecutive_hyphen_limit = OxmlElement("w:consecutiveHyphenLimit")
    consecutive_hyphen_limit.set(qn("w:val"), "2")
    sectPr.append(consecutive_hyphen_limit)


def set_compatibility_options(doc):
    """Set compatibility options to optimize spacing and justification for research paper quality with equal line lengths."""
    compat = doc.settings.element.find(qn("w:compat"))
    if compat is None:
        doc.settings.element.append(OxmlElement("w:compat"))
        compat = doc.settings.element.find(qn("w:compat"))

    # Critical options for professional justification with equal line lengths

    # Use Word 2010+ justification algorithm for better spacing
    option1 = OxmlElement("w:useWord2010TableStyleRules")
    option1.set(qn("w:val"), "1")
    compat.append(option1)

    # Enable better line breaking for justified text
    option2 = OxmlElement("w:doNotBreakWrappedTables")
    option2.set(qn("w:val"), "1")
    compat.append(option2)

    # Use consistent font metrics for better spacing
    option3 = OxmlElement("w:useWord97LineBreakRules")
    option3.set(qn("w:val"), "0")  # Disable old line break rules
    compat.append(option3)

    # Enable advanced justification for equal line lengths
    option4 = OxmlElement("w:doNotExpandShiftReturn")
    option4.set(qn("w:val"), "1")  # Better line break handling
    compat.append(option4)

    # Force consistent character spacing
    option5 = OxmlElement("w:doNotUseEastAsianBreakRules")
    option5.set(qn("w:val"), "1")  # Use Western justification rules
    compat.append(option5)

    # Enable text compression for equal line fitting
    option6 = OxmlElement("w:allowSpaceOfSameStyleInTable")
    option6.set(qn("w:val"), "1")  # Better spacing in justified text
    compat.append(option6)

    # Prevent Word from expanding spaces for justification
    option2 = OxmlElement("w:doNotExpandShiftReturn")
    option2.set(qn("w:val"), "1")
    compat.append(option2)

    # Use consistent character spacing
    option3 = OxmlElement("w:useSingleBorderforContiguousCells")
    option3.set(qn("w:val"), "1")
    compat.append(option3)

    # Force exact spacing calculations
    option4 = OxmlElement("w:spacingInWholePoints")
    option4.set(qn("w:val"), "1")
    compat.append(option4)

    # Prevent auto spacing adjustments
    option5 = OxmlElement("w:doNotUseHTMLParagraphAutoSpacing")
    option5.set(qn("w:val"), "1")
    compat.append(option5)

    # Use legacy justification method (more precise)
    option6 = OxmlElement("w:useWord97LineBreakRules")
    option6.set(qn("w:val"), "1")
    compat.append(option6)

    # Disable automatic kerning adjustments
    option7 = OxmlElement("w:doNotAutoCompressPictures")
    option7.set(qn("w:val"), "1")
    compat.append(option7)

    # Force consistent text metrics
    option8 = OxmlElement("w:useNormalStyleForList")
    option8.set(qn("w:val"), "1")
    compat.append(option8)

    # Prevent text compression/expansion
    option9 = OxmlElement("w:doNotPromoteQF")
    option9.set(qn("w:val"), "1")
    compat.append(option9)

    # Use exact font metrics
    option10 = OxmlElement("w:useAltKinsokuLineBreakRules")
    option10.set(qn("w:val"), "0")
    compat.append(option10)


def generate_ieee_document(form_data):
    """Generate IEEE-formatted Word document with EXACT LaTeX PDF formatting via OpenXML."""
    doc = Document()

    # Apply EXACT IEEE LaTeX PDF specifications
    set_document_defaults(doc)

    # Configure first section for single-column title and authors (IEEE LaTeX standard)
    section = doc.sections[0]
    section.left_margin = Inches(0.75)  # EXACT IEEE LaTeX: 0.75" margins
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Add title and authors in single-column layout (EXACT IEEE LaTeX standard)
    add_title(doc, form_data.get("title", ""))
    add_authors(doc, form_data.get("authors", []))

    # Setup TWO-COLUMN LAYOUT for body content (EXACT IEEE LaTeX specifications)
    setup_two_column_layout(doc)

    # Add abstract and keywords in two-column layout with EXACT IEEE LaTeX formatting
    add_abstract(doc, form_data.get("abstract", ""))
    add_keywords(doc, form_data.get("keywords", ""))

    # Add sections with EXACT IEEE LaTeX formatting
    for idx, section_data in enumerate(form_data.get("sections", []), 1):
        add_section(doc, section_data, idx, is_first_section=(idx == 1))

    # NOTE: Standalone tables and figures are now processed through the content blocks system
    # This ensures proper ordering and prevents duplication
    print("Standalone tables and figures will be processed through content blocks for proper ordering", file=sys.stderr)

    # Add references with EXACT IEEE LaTeX formatting
    add_references(doc, form_data.get("references", []))

    # Apply final IEEE LaTeX compatibility settings
    enable_auto_hyphenation(doc)
    set_compatibility_options(doc)

    # Generate final document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_document_model(form_data):
    """Build structured document model with exact OpenXML-level formatting metadata"""
    
    # Extract and sanitize data
    title = sanitize_text(form_data.get("title", "Untitled Document"))
    authors = form_data.get("authors", [])
    abstract = sanitize_text(form_data.get("abstract", ""))
    keywords = sanitize_text(form_data.get("keywords", ""))
    sections = form_data.get("sections", [])
    references = form_data.get("references", [])
    
    # Build structured model with exact formatting specifications
    model = {
        "title": {
            "text": title,
            "font_size": "24pt",  # IEEE_CONFIG["font_size_title"] = Pt(24)
            "font_weight": "bold",
            "text_align": "center",
            "margin_bottom": "12pt",  # 240 twips = 12pt
            "font_family": "Times New Roman"
        },
        "authors": [],
        "abstract": {
            "text": abstract,
            "font_size": "9pt",  # IEEE_CONFIG["font_size_abstract"] = Pt(9)
            "font_weight": "bold",
            "text_align": "justify",
            "margin_bottom": "6pt",  # 120 twips = 6pt
            "prefix": "Abstract—"
        } if abstract else None,
        "keywords": {
            "text": keywords,
            "font_size": "9pt",
            "font_weight": "bold", 
            "text_align": "justify",
            "margin_bottom": "12pt",  # 240 twips = 12pt
            "prefix": "Index Terms—"
        } if keywords else None,
        "sections": [],
        "references": {
            "items": [],
            "font_size": "9pt",  # IEEE_CONFIG["font_size_reference"] = Pt(9)
            "hanging_indent": "0.25in",  # 360 twips = 0.25in
            "text_align": "justify"
        } if references else None,
        "page_config": {
            "margins": "0.75in",  # 1080 twips = 0.75in
            "font_family": "Times New Roman",
            "body_font_size": "10pt",  # IEEE_CONFIG["font_size_body"] = Pt(10)
            "line_height": "12pt",  # 240 twips = 12pt
            "column_count": 2,
            "column_gap": "0.25in",  # 360 twips = 0.25in
            "column_width": "3.3125in"  # 4770 twips = 3.3125in
        }
    }
    
    # Process authors with exact 3-column layout
    if authors:
        authors_per_row = 3
        for row_start in range(0, len(authors), authors_per_row):
            row_end = min(row_start + authors_per_row, len(authors))
            row_authors = authors[row_start:row_end]
            
            author_row = {
                "type": "author_row",
                "authors": [],
                "columns": 3,
                "font_size": "10pt",
                "text_align": "center",
                "margin_bottom": "10pt" if row_end < len(authors) else "20pt"
            }
            
            for author in row_authors:
                author_data = {
                    "name": sanitize_text(author.get("name", "")),
                    "name_font_weight": "bold",
                    "name_font_size": "10pt",
                    "affiliation_font_style": "italic",
                    "affiliation_font_size": "10pt",
                    "email_font_size": "9pt",
                    "fields": []
                }
                
                # Add structured fields in IEEE order
                fields = ["department", "organization", "university", "institution", "city", "state", "country"]
                for field in fields:
                    if author.get(field):
                        author_data["fields"].append({
                            "type": "affiliation",
                            "text": sanitize_text(author[field])
                        })
                
                # Add email
                if author.get("email"):
                    author_data["fields"].append({
                        "type": "email", 
                        "text": sanitize_text(author["email"])
                    })
                
                # Fallback to affiliation field
                if not author_data["fields"] and author.get("affiliation"):
                    for line in author["affiliation"].strip().split("\n"):
                        line = line.strip()
                        if line:
                            author_data["fields"].append({
                                "type": "affiliation",
                                "text": sanitize_text(line)
                            })
                
                author_row["authors"].append(author_data)
            
            model["authors"].append(author_row)
    
    # Process standalone tables and figures first (from table-form.tsx and figure-form.tsx)
    standalone_tables = form_data.get("tables", [])
    standalone_figures = form_data.get("figures", [])
    
    # Add standalone tables and figures to the first section or create a new section
    if standalone_tables or standalone_figures:
        # Find or create a section for standalone content
        if not sections:
            sections.append({
                "title": "Content",
                "contentBlocks": []
            })
        
        # Add standalone tables as content blocks with proper ordering
        for table in standalone_tables:
            table_block = {
                "type": "table",
                "tableType": table.get("type", table.get("tableType", "interactive")),
                "tableName": table.get("tableName", ""),
                "caption": table.get("caption", ""),
                "size": table.get("size", "medium"),
                "order": table.get("order", 999)  # Use table's order or default to end
            }
            
            if table_block["tableType"] == "interactive":
                table_block["headers"] = table.get("headers", [])
                table_block["tableData"] = table.get("tableData", [])
            elif table_block["tableType"] == "image":
                table_block["data"] = table.get("data", "")
                table_block["originalName"] = table.get("originalName", "")
                table_block["mimeType"] = table.get("mimeType", "")
            elif table_block["tableType"] == "latex":
                table_block["latexCode"] = table.get("latexCode", "")
            
            sections[0].setdefault("contentBlocks", []).append(table_block)
        
        # Add standalone figures as content blocks with proper ordering
        for figure in standalone_figures:
            figure_block = {
                "type": "image",
                "data": figure.get("data", ""),
                "caption": figure.get("caption", ""),
                "size": figure.get("size", "medium"),
                "originalName": figure.get("originalName", ""),
                "mimeType": figure.get("mimeType", ""),
                "order": figure.get("order", 999)  # Use figure's order or default to end
            }
            sections[0].setdefault("contentBlocks", []).append(figure_block)

    # Process sections with content blocks
    for section_idx, section in enumerate(sections, 1):
        section_title = sanitize_text(section.get("title", ""))
        section_data = {
            "number": section_idx,
            "title": section_title,
            "font_size": "10pt",
            "font_weight": "bold",
            "text_align": "center",
            "text_transform": "uppercase",
            "margin_top": "12pt",  # 240 twips = 12pt
            "margin_bottom": "0pt",
            "content_blocks": []
        }
        
        # Process content blocks (including converted standalone tables/figures)
        content_blocks = section.get("contentBlocks", [])
        
        # CRITICAL FIX: Sort content blocks by their order field to ensure correct sequence
        content_blocks = sorted(content_blocks, key=lambda x: x.get("order", 999))
        
        table_count = 0
        img_count = 0
        
        for block in content_blocks:
            block_type = block.get("type", "text")
            
            if block_type == "text" and block.get("content"):
                section_data["content_blocks"].append({
                    "type": "paragraph",
                    "text": sanitize_text(block["content"]),
                    "font_size": "10pt",
                    "text_align": "justify",
                    "margin_bottom": "12pt",
                    "line_height": "12pt"
                })
            
            elif block_type == "table":
                table_count += 1
                table_type = block.get("tableType", "interactive")
                
                if table_type == "interactive":
                    headers = block.get("headers", [])
                    rows_data = block.get("tableData", [])
                    
                    if headers and rows_data:
                        table_data = {
                            "type": "table",
                            "number": f"{section_idx}.{table_count}",
                            "headers": [sanitize_text(str(h)) for h in headers],
                            "rows": [[sanitize_text(str(cell)) for cell in row] for row in rows_data],
                            "font_size": "9pt",
                            "border": "1px solid black",
                            "margin": "12pt 0",
                            "width": "100%",
                            "caption": {
                                "text": sanitize_text(block.get("caption", block.get("tableName", ""))),
                                "font_size": "9pt",
                                "font_weight": "bold",
                                "text_align": "center",
                                "margin": "6pt 0 12pt 0"
                            }
                        }
                        section_data["content_blocks"].append(table_data)
                
                elif table_type == "image" and block.get("data"):
                    table_count += 1
                    image_data = block["data"]
                    if "," in image_data:
                        image_data = image_data.split(",")[1]
                    
                    size_mapping = {
                        "very-small": "1.5in",
                        "small": "2.0in", 
                        "medium": "2.5in",
                        "large": "3.3125in"
                    }
                    
                    table_image_data = {
                        "type": "table_image",
                        "number": f"{section_idx}.{table_count}",
                        "data": image_data,
                        "width": size_mapping.get(block.get("size", "medium"), "2.5in"),
                        "text_align": "center",
                        "margin": "12pt 0",
                        "caption": {
                            "text": sanitize_text(block.get("caption", block.get("tableName", ""))),
                            "font_size": "9pt",
                            "font_weight": "bold",
                            "text_align": "center",
                            "text_transform": "uppercase",
                            "margin": "6pt 0 12pt 0",
                            "prefix": f"TABLE {section_idx}.{table_count}: "
                        }
                    }
                    section_data["content_blocks"].append(table_image_data)
            
            elif block_type == "image" and block.get("data") and block.get("caption"):
                img_count += 1
                image_data = block["data"]
                if "," in image_data:
                    image_data = image_data.split(",")[1]
                
                size_mapping = {
                    "very-small": "1.5in",
                    "small": "2.0in",
                    "medium": "2.5in", 
                    "large": "3.3125in"
                }
                
                image_block_data = {
                    "type": "figure",
                    "number": f"{section_idx}.{img_count}",
                    "data": image_data,
                    "width": size_mapping.get(block.get("size", "medium"), "2.5in"),
                    "text_align": "center",
                    "margin": "12pt 0",
                    "caption": {
                        "text": sanitize_text(block["caption"]),
                        "font_size": "9pt",
                        "font_weight": "bold",
                        "text_align": "center",
                        "text_transform": "uppercase",
                        "margin": "6pt 0 12pt 0",
                        "prefix": f"FIG. {section_idx}.{img_count}: "
                    }
                }
                section_data["content_blocks"].append(image_block_data)
            
            elif block_type == "equation" and block.get("content"):
                equation_data = {
                    "type": "equation",
                    "content": sanitize_text(block["content"]),
                    "number": block.get("equationNumber", ""),
                    "font_size": "10pt",
                    "text_align": "center",
                    "margin": "12pt 0",
                    "font_style": "italic"
                }
                section_data["content_blocks"].append(equation_data)
        
        model["sections"].append(section_data)
    
    # Process references
    if references:
        ref_items = []
        for i, ref in enumerate(references, 1):
            ref_text = sanitize_text(ref.get("text", "")) if isinstance(ref, dict) else sanitize_text(str(ref))
            if ref_text:
                ref_items.append({
                    "number": i,
                    "text": ref_text,
                    "margin": "3pt 0"
                })
        model["references"]["items"] = ref_items
    
    return model


def render_to_html(model):
    """Render document model to pixel-perfect HTML matching OpenXML formatting"""
    
    page_config = model["page_config"]
    
    # Build CSS with exact OpenXML measurements converted to CSS units
    css = f"""
        /* EXACT IEEE LaTeX PDF SPECIFICATIONS - Pixel Perfect OpenXML Conversion */
        
        @page {{
            size: letter;
            margin: {page_config["margins"]};
        }}
        
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            font-family: '{page_config["font_family"]}', Times, serif;
            font-size: {page_config["body_font_size"]};
            line-height: {page_config["line_height"]};
            color: black;
            background: white;
            
            /* PERFECT JUSTIFICATION - Match Word's OpenXML justification */
            text-align: justify;
            text-justify: distribute;
            hyphens: auto;
            -webkit-hyphens: auto;
            -moz-hyphens: auto;
            -ms-hyphens: auto;
            
            /* Enhanced typography matching OpenXML settings */
            text-rendering: optimizeLegibility;
            font-variant-ligatures: common-ligatures;
            font-feature-settings: "liga" 1, "kern" 1;
            
            /* Fine-tune spacing for perfect justification */
            word-spacing: 0.05em;
            letter-spacing: -0.01em;
            
            /* Prevent orphans and widows */
            orphans: 2;
            widows: 2;
        }}
        
        /* TITLE - Exact match to OpenXML title formatting */
        .ieee-title {{
            font-size: {model["title"]["font_size"]};
            font-weight: {model["title"]["font_weight"]};
            text-align: {model["title"]["text_align"]};
            margin-bottom: {model["title"]["margin_bottom"]};
            line-height: 1.3;
            page-break-after: avoid;
            letter-spacing: 0;
            word-spacing: 0;
        }}
        
        /* AUTHORS - CSS Grid matching OpenXML table layout */
        .ieee-authors-container {{
            margin: 15px 0 20px 0;
            page-break-after: avoid;
        }}
        
        .ieee-authors-row {{
            display: grid;
            grid-template-columns: 1fr 1fr 1fr;
            gap: 0.25in;
            margin-bottom: 10px;
            text-align: center;
        }}
        
        .ieee-author {{
            font-size: 10pt;
            line-height: 1.2;
        }}
        
        .author-name {{
            font-weight: bold;
            margin-bottom: 3px;
        }}
        
        .author-affiliation {{
            font-style: italic;
            margin-bottom: 2px;
        }}
        
        .author-email {{
            font-size: 9pt;
            margin-top: 2px;
        }}
        
        /* ABSTRACT and KEYWORDS - Exact OpenXML formatting */
        .ieee-abstract, .ieee-keywords {{
            font-size: 9pt;
            font-weight: bold;
            margin: 15px 0;
            text-align: justify;
            text-justify: distribute;
            hyphens: auto;
            break-inside: avoid;
            word-spacing: 0.05em;
            letter-spacing: -0.01em;
        }}
        
        /* TWO-COLUMN LAYOUT - Start after abstract/keywords */
        .ieee-body {{
            columns: {page_config["column_count"]};
            column-gap: {page_config["column_gap"]};
            column-fill: balance;
            column-rule: none;
        }}
        
        /* SECTION HEADINGS - Exact OpenXML formatting */
        .ieee-heading {{
            font-size: 10pt;
            font-weight: bold;
            text-align: center;
            text-transform: uppercase;
            margin: 12pt 0 0 0;
            page-break-after: avoid;
            break-after: avoid;
            letter-spacing: 0;
            word-spacing: 0;
        }}
        
        /* PARAGRAPHS - Perfect justification matching OpenXML */
        .ieee-paragraph {{
            font-size: 10pt;
            margin: 0 0 12pt 0;
            text-align: justify;
            text-justify: distribute;
            hyphens: auto;
            letter-spacing: -0.01em;
            word-spacing: 0.05em;
            orphans: 2;
            widows: 2;
            line-height: 12pt;
        }}
        
        /* TABLES - Exact IEEE formatting */
        .ieee-table-container {{
            margin: 12pt 0;
            break-inside: avoid;
            page-break-inside: avoid;
        }}
        
        .ieee-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 9pt;
            margin: 6pt auto;
            border: 1px solid black;
        }}
        
        .ieee-table-header {{
            border: 1px solid black;
            padding: 4pt 6pt;
            text-align: center;
            font-weight: bold;
            background-color: #f5f5f5;
            vertical-align: middle;
        }}
        
        .ieee-table-cell {{
            border: 1px solid black;
            padding: 4pt 6pt;
            text-align: left;
            vertical-align: top;
        }}
        
        .ieee-table-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6pt 0 12pt 0;
            break-before: avoid;
        }}
        
        /* IMAGES - Exact sizing and positioning */
        .ieee-image-container {{
            text-align: center;
            margin: 12pt 0;
            break-inside: avoid;
            page-break-inside: avoid;
        }}
        
        .ieee-image {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 0 auto;
        }}
        
        .ieee-figure-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6pt 0 12pt 0;
            break-before: avoid;
        }}
        
        /* EQUATIONS - IEEE formatting */
        .ieee-equation-container {{
            margin: 12pt 0;
            text-align: center;
            break-inside: avoid;
            page-break-inside: avoid;
        }}
        
        .ieee-equation {{
            font-size: 10pt;
            font-style: italic;
            display: inline-block;
            padding: 6pt 12pt;
            background-color: #f9f9f9;
            border: 1px solid #ddd;
            border-radius: 4px;
        }}
        
        /* REFERENCES - Exact hanging indent */
        .ieee-reference {{
            font-size: 9pt;
            margin: 3pt 0;
            padding-left: 0.25in;
            text-indent: -0.25in;
            text-align: justify;
            text-justify: distribute;
            hyphens: auto;
            letter-spacing: -0.02em;
            word-spacing: 0.05em;
        }}
        
        /* PAGE BREAKS */
        .page-break {{
            page-break-before: always;
            break-before: page;
        }}
        
        .keep-together {{
            break-inside: avoid;
            page-break-inside: avoid;
        }}
        
        /* PRINT OPTIMIZATIONS */
        @media print {{
            body {{
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
            
            .ieee-table {{
                border-collapse: collapse !important;
            }}
            
            .ieee-table-header,
            .ieee-table-cell {{
                border: 1px solid black !important;
            }}
        }}
    """
    
    # Build HTML structure
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{model["title"]["text"]}</title>
    <style>{css}</style>
</head>
<body>
    <div class="ieee-title">{model["title"]["text"]}</div>
"""
    
    # Add authors
    if model["authors"]:
        html += '<div class="ieee-authors-container">'
        for author_row in model["authors"]:
            html += '<div class="ieee-authors-row">'
            
            # Always create 3 columns for consistent layout
            for i in range(3):
                if i < len(author_row["authors"]):
                    author = author_row["authors"][i]
                    html += '<div class="ieee-author">'
                    html += f'<div class="author-name">{author["name"]}</div>'
                    
                    for field in author["fields"]:
                        if field["type"] == "email":
                            html += f'<div class="author-email">{field["text"]}</div>'
                        else:
                            html += f'<div class="author-affiliation">{field["text"]}</div>'
                    
                    html += '</div>'
                else:
                    html += '<div class="ieee-author"></div>'  # Empty column
            
            html += '</div>'
        html += '</div>'
    
    # Add abstract and keywords (single column)
    if model["abstract"]:
        html += f'<div class="ieee-abstract"><strong>{model["abstract"]["prefix"]}</strong>{model["abstract"]["text"]}</div>'
    
    if model["keywords"]:
        html += f'<div class="ieee-keywords"><strong>{model["keywords"]["prefix"]}</strong>{model["keywords"]["text"]}</div>'
    
    # Start two-column layout for body content
    html += '<div class="ieee-body">'
    
    # Add sections
    for section in model["sections"]:
        if section["title"]:
            html += f'<div class="ieee-heading">{section["number"]}. {section["title"]}</div>'
        
        for block in section["content_blocks"]:
            if block["type"] == "paragraph":
                html += f'<div class="ieee-paragraph">{block["text"]}</div>'
            
            elif block["type"] == "table":
                html += '<div class="ieee-table-container">'
                html += '<table class="ieee-table">'
                
                # Headers
                html += '<thead><tr>'
                for header in block["headers"]:
                    html += f'<th class="ieee-table-header">{header}</th>'
                html += '</tr></thead>'
                
                # Rows
                html += '<tbody>'
                for row in block["rows"]:
                    html += '<tr>'
                    for cell in row:
                        html += f'<td class="ieee-table-cell">{cell}</td>'
                    html += '</tr>'
                html += '</tbody>'
                
                html += '</table>'
                
                if block["caption"]["text"]:
                    html += f'<div class="ieee-table-caption">TABLE {block["number"]}: {block["caption"]["text"].upper()}</div>'
                
                html += '</div>'
            
            elif block["type"] == "table_image":
                html += '<div class="ieee-image-container">'
                if block["caption"]["text"]:
                    html += f'<div class="ieee-table-caption">{block["caption"]["prefix"]}{block["caption"]["text"].upper()}</div>'
                html += f'<img src="data:image/png;base64,{block["data"]}" class="ieee-image" style="width: {block["width"]};" alt="Table {block["number"]}" />'
                html += '</div>'
            
            elif block["type"] == "figure":
                html += '<div class="ieee-image-container">'
                html += f'<img src="data:image/png;base64,{block["data"]}" class="ieee-image" style="width: {block["width"]};" alt="Figure {block["number"]}" />'
                if block["caption"]["text"]:
                    html += f'<div class="ieee-figure-caption">{block["caption"]["prefix"]}{block["caption"]["text"].upper()}</div>'
                html += '</div>'
            
            elif block["type"] == "equation":
                html += '<div class="ieee-equation-container">'
                equation_content = block["content"]
                equation_number = block.get("number", "")
                if equation_number:
                    html += f'<div class="ieee-equation">({equation_number}) {equation_content}</div>'
                else:
                    html += f'<div class="ieee-equation">{equation_content}</div>'
                html += '</div>'
    
    # Add references
    if model["references"]:
        html += '<div class="ieee-heading">REFERENCES</div>'
        for ref in model["references"]["items"]:
            html += f'<div class="ieee-reference">[{ref["number"]}] {ref["text"]}</div>'
    
    html += '</div>'  # Close ieee-body
    html += '</body></html>'
    
    return html


def generate_ieee_html_preview(form_data):
    """Generate HTML preview using unified rendering system - 100% identical to PDF"""
    model = build_document_model(form_data)
    html = render_to_html(model)
    
    # Add preview note for live preview
    preview_note = '''
    <div style="background: #e8f4fd; border: 1px solid #bee5eb; padding: 12px; margin: 20px 0; font-size: 9pt; color: #0c5460; text-align: center; border-radius: 4px;">
        📄 IEEE Live Preview - This is exactly what your PDF will look like
    </div>
    '''
    
    # Insert preview note after body tag
    html = html.replace('<body>', f'<body>{preview_note}', 1)
    
    return html

    # Create HTML with IDENTICAL justification for both preview and PDF
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>{title}</title>
        <style>
            @page {{
                size: letter;
                margin: 0.75in;
            }}

            body {{
                font-family: 'Times New Roman', serif;
                font-size: 10pt;
                line-height: 1.2;
                margin: 0.75in;
                background: white;
                color: black;

                /* PERFECT JUSTIFICATION - Works in both browsers and WeasyPrint */
                text-align: justify;
                text-justify: inter-word;
                hyphens: auto;
                -webkit-hyphens: auto;
                -moz-hyphens: auto;
                -ms-hyphens: auto;

                /* Enhanced typography */
                text-rendering: optimizeLegibility;
                font-variant-ligatures: common-ligatures;
                font-feature-settings: "liga" 1, "kern" 1;

                /* Fine-tune spacing for better justification */
                word-spacing: 0.05em;
                letter-spacing: -0.01em;
            }}

            .ieee-title {{
                font-size: 24pt;
                font-weight: bold;
                text-align: center;
                margin: 20px 0;
                line-height: 1.3;
            }}

            .ieee-authors {{
                font-size: 10pt;
                text-align: center;
                margin: 15px 0;
            }}

            .ieee-section {{
                margin: 15px 0;
                text-align: justify;
                text-justify: inter-word;
                hyphens: auto;
                -webkit-hyphens: auto;
                -moz-hyphens: auto;
                -ms-hyphens: auto;
                word-spacing: 0.05em;
                letter-spacing: -0.01em;
            }}

            .ieee-abstract-title, .ieee-keywords-title {{
                font-weight: bold;
                display: inline;
            }}

            .ieee-heading {{
                font-weight: bold;
                margin: 15px 0 5px 0;
                text-transform: uppercase;
                font-size: 10pt;
                text-align: center;
            }}

            .ieee-reference {{
                margin: 3px 0;
                padding-left: 15px;
                text-indent: -15px;
                font-size: 9pt;
                text-align: justify;
                text-justify: inter-word;
                hyphens: auto;
                word-spacing: 0.05em;
                letter-spacing: -0.01em;
            }}

            .preview-note {{
                background: #e8f4fd;
                border: 1px solid #bee5eb;
                padding: 12px;
                margin: 20px 0;
                font-size: 9pt;
                color: #0c5460;
                text-align: center;
                border-radius: 4px;
            }}

            .content-block {{
                margin: 10px 0;
                text-align: justify;
                text-justify: inter-word;
                hyphens: auto;
                -webkit-hyphens: auto;
                -moz-hyphens: auto;
                -ms-hyphens: auto;
                word-spacing: 0.05em;
                letter-spacing: -0.01em;
            }}

            .figure-caption {{
                font-size: 9pt;
                text-align: center;
                margin: 10px 0;
                font-style: italic;
            }}

            /* UNIVERSAL JUSTIFICATION - Apply to ALL text elements */
            p, div:not(.ieee-title):not(.ieee-authors):not(.ieee-heading):not(.figure-caption):not(.preview-note) {{
                text-align: justify;
                text-justify: inter-word;
                hyphens: auto;
                -webkit-hyphens: auto;
                -moz-hyphens: auto;
                -ms-hyphens: auto;
                word-spacing: 0.05em;
                letter-spacing: -0.01em;
            }}

            /* WeasyPrint specific enhancements */
            @media print {{
                body, p, div, .ieee-section, .content-block, .ieee-reference {{
                    text-align: justify !important;
                    hyphens: auto !important;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="preview-note">
            📄 IEEE Live Preview - This is exactly what your PDF will look like
        </div>

        <div class="ieee-title">{title}</div>
        <div class="ieee-authors">{authors_html}</div>
    """

    # Add abstract
    if abstract:
        html += f"""
        <div class="ieee-section">
            <span class="ieee-abstract-title">Abstract—</span>{abstract}
        </div>
        """

    # Add keywords
    if keywords:
        html += f"""
        <div class="ieee-section">
            <span class="ieee-keywords-title">Index Terms—</span>{keywords}
        </div>
        """

    # Add sections with content blocks
    for i, section in enumerate(sections, 1):
        section_title = sanitize_text(section.get("title", ""))
        if section_title:
            html += f"""
            <div class="ieee-heading">{i}. {section_title}</div>
            """

            # Process content blocks within the section
            content_blocks = section.get("contentBlocks", [])
            for block in content_blocks:
                block_type = block.get("type", "text")
                block_content = sanitize_text(block.get("content", ""))

                if block_type == "text" and block_content:
                    html += f'<div class="content-block">{block_content}</div>'
                elif block_type == "figure" and block_content:
                    html += f'<div class="content-block figure-caption">Fig. {i}. {block_content}</div>'
                elif block_type == "equation" and block_content:
                    html += f'<div class="content-block" style="text-align: center; margin: 15px 0;">{block_content}</div>'

    # Add references
    if references:
        html += '<div class="ieee-heading">References</div>'
        for i, ref in enumerate(references, 1):
            ref_text = sanitize_text(ref.get("text", ""))
            if ref_text:
                html += f'<div class="ieee-reference">[{i}] {ref_text}</div>'

    html += """
        <div class="preview-note">
            ✨ PDF download will have identical formatting to this preview
        </div>
    </body>
    </html>
    """

    return html


def generate_ieee_master_html(form_data):
    """Generate MASTER HTML with pixel-perfect IEEE formatting - used by both DOCX and PDF outputs"""

    # Extract document data
    title = sanitize_text(form_data.get("title", "Untitled Document"))
    authors = form_data.get("authors", [])
    abstract = sanitize_text(form_data.get("abstract", ""))
    keywords = sanitize_text(form_data.get("keywords", ""))
    sections = form_data.get("sections", [])
    references = form_data.get("references", [])

    # Format authors with CSS Grid for exact 3-column layout (IEEE standard)
    authors_html = ""
    if authors:
        authors_html = '<div class="ieee-authors-container">'

        # Process authors in groups of 3 (IEEE standard)
        authors_per_row = 3
        total_authors = len(authors)

        for row_start in range(0, total_authors, authors_per_row):
            row_end = min(row_start + authors_per_row, total_authors)
            row_authors = authors[row_start:row_end]

            authors_html += '<div class="ieee-authors-row">'

            for author in row_authors:
                author_name = sanitize_text(author.get("name", ""))
                author_html = f'<div class="ieee-author"><div class="author-name">{author_name}</div>'

                # Add structured affiliation fields in IEEE order
                fields = [
                    "department",
                    "organization",
                    "university",
                    "institution",
                    "city",
                    "state",
                    "country",
                ]
                for field in fields:
                    if author.get(field):
                        author_html += f'<div class="author-affiliation">{sanitize_text(author[field])}</div>'

                # Add email
                if author.get("email"):
                    author_html += f'<div class="author-email">{sanitize_text(author["email"])}</div>'

                # Fallback to affiliation field if structured fields not available
                if not any(author.get(field) for field in fields) and author.get(
                    "affiliation"
                ):
                    affiliation_lines = author["affiliation"].strip().split("\n")
                    for line in affiliation_lines:
                        line = line.strip()
                        if line:
                            author_html += f'<div class="author-affiliation">{sanitize_text(line)}</div>'

                author_html += "</div>"
                authors_html += author_html

            # Fill remaining columns if less than 3 authors in this row
            remaining_cols = authors_per_row - len(row_authors)
            for _ in range(remaining_cols):
                authors_html += (
                    '<div class="ieee-author"></div>'  # Empty column for grid alignment
                )

            authors_html += "</div>"

        authors_html += "</div>"

    # Process sections with content blocks (tables and images)
    sections_html = ""
    for section_idx, section in enumerate(sections, 1):
        section_title = sanitize_text(section.get("title", ""))
        if section_title:
            sections_html += f'<div class="ieee-heading">{section_idx}. {section_title.upper()}</div>'

        # Process content blocks
        content_blocks = section.get("contentBlocks", [])
        table_count = 0
        img_count = 0

        for block in content_blocks:
            block_type = block.get("type", "text")

            if block_type == "text" and block.get("content"):
                content = sanitize_text(block["content"])
                sections_html += f'<div class="ieee-paragraph">{content}</div>'

            elif block_type == "table":
                table_count += 1
                table_type = block.get("tableType", "interactive")

                if table_type == "interactive":
                    headers = block.get("headers", [])
                    rows_data = block.get("tableData", [])

                    if headers and rows_data:
                        sections_html += '<div class="ieee-table-container">'
                        sections_html += '<table class="ieee-table">'

                        # Header row
                        sections_html += "<thead><tr>"
                        for header in headers:
                            sections_html += f'<th class="ieee-table-header">{sanitize_text(str(header))}</th>'
                        sections_html += "</tr></thead>"

                        # Data rows
                        sections_html += "<tbody>"
                        for row_data in rows_data:
                            sections_html += "<tr>"
                            for cell_data in row_data:
                                sections_html += f'<td class="ieee-table-cell">{sanitize_text(str(cell_data))}</td>'
                            sections_html += "</tr>"
                        sections_html += "</tbody>"

                        sections_html += "</table>"

                        # Table caption
                        # Fix table caption duplication
                        caption_text = block.get("caption", "").strip()
                        table_name = block.get("tableName", "").strip()

                        if caption_text and table_name:
                            if table_name.lower() in caption_text.lower():
                                final_caption = caption_text
                            elif caption_text.lower() in table_name.lower():
                                final_caption = table_name
                            else:
                                final_caption = caption_text
                        else:
                            final_caption = caption_text or table_name

                        if final_caption:
                            sections_html += f'<div class="ieee-table-caption">TABLE {section_idx}.{table_count}: {sanitize_text(final_caption).upper()}</div>'

                        sections_html += "</div>"

                elif table_type == "image" and block.get("data"):
                    # Handle image tables - ENSURE PROPER DISPLAY IN WORD
                    image_data = block["data"]
                    if "," in image_data:
                        image_data = image_data.split(",")[1]

                    # Get table name and caption
                    table_name = block.get(
                        "tableName",
                        block.get("caption", f"Table {section_idx}.{table_count}"),
                    )
                    caption = block.get("caption", block.get("tableName", ""))

                    size_class = f"ieee-image-{block.get('size', 'medium')}"
                    sections_html += f'<div class="ieee-image-container">'

                    # Add table name BEFORE image for Word compatibility
                    if table_name:
                        sections_html += f'<div class="ieee-table-name">TABLE {section_idx}.{table_count}: {sanitize_text(table_name).upper()}</div>'

                    # Image with proper alt text including table name
                    alt_text = (
                        f"Table {section_idx}.{table_count}: {table_name}"
                        if table_name
                        else f"Table {section_idx}.{table_count}"
                    )
                    sections_html += f'<img src="data:image/png;base64,{image_data}" class="ieee-image {size_class}" alt="{alt_text}" title="{alt_text}" />'

                    # Caption AFTER image
                    if caption and caption != table_name:
                        sections_html += f'<div class="ieee-table-caption">{sanitize_text(caption)}</div>'

                    sections_html += "</div>"

            elif block_type == "image" and block.get("data") and block.get("caption"):
                img_count += 1
                image_data = block["data"]
                if "," in image_data:
                    image_data = image_data.split(",")[1]

                size_class = f"ieee-image-{block.get('size', 'medium')}"
                sections_html += f'<div class="ieee-image-container">'
                sections_html += f'<img src="data:image/png;base64,{image_data}" class="ieee-image {size_class}" alt="Figure {section_idx}.{img_count}" />'
                sections_html += f'<div class="ieee-figure-caption">FIG. {section_idx}.{img_count}: {sanitize_text(block["caption"]).upper()}</div>'
                sections_html += "</div>"

    # Process references
    references_html = ""
    if references:
        references_html = '<div class="ieee-heading">REFERENCES</div>'
        for i, ref in enumerate(references, 1):
            ref_text = (
                sanitize_text(ref.get("text", ""))
                if isinstance(ref, dict)
                else sanitize_text(str(ref))
            )
            if ref_text:
                references_html += f'<div class="ieee-reference">[{i}] {ref_text}</div>'

    # Create MASTER HTML with EXACT IEEE CSS - identical for both DOCX and PDF
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        /* EXACT IEEE LaTeX PDF SPECIFICATIONS - PIXEL PERFECT */

        @page {{
            size: letter;
            margin: 0.75in;
            @bottom-center {{
                content: counter(page);
                font-family: 'Times New Roman', serif;
                font-size: 10pt;
            }}
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}

        body {{
            font-family: 'Times New Roman', serif;
            font-size: 10pt;
            line-height: 1.2;
            color: black;
            background: white;

            /* ULTRA-AGGRESSIVE PERFECT JUSTIFICATION - Force LaTeX quality */
            text-align: justify !important;
            text-justify: distribute !important;
            text-align-last: justify !important;
            hyphens: auto !important;
            -webkit-hyphens: auto !important;
            -moz-hyphens: auto !important;
            -ms-hyphens: auto !important;

            /* ULTRA-AGGRESSIVE character spacing for perfect line endings */
            letter-spacing: 0.02em !important;
            word-spacing: 0.12em !important;

            /* WeasyPrint specific justification */
            -weasy-text-align-last: justify !important;
            -weasy-text-justify: distribute !important;
            -weasy-hyphens: auto !important;

            /* Typography controls */
            text-rendering: optimizeLegibility;
            font-variant-ligatures: common-ligatures;
            font-feature-settings: "liga" 1, "kern" 1;

            /* Prevent orphans and widows */
            orphans: 2;
            widows: 2;
        }}

        /* TITLE - 24pt bold centered */
        .ieee-title {{
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin: 0 0 20px 0;
            line-height: 1.3;
            page-break-after: avoid;
            letter-spacing: 0;
            word-spacing: 0;
        }}

        /* AUTHORS - CSS Grid for exact 3-column layout */
        .ieee-authors-container {{
            margin: 15px 0 20px 0;
            page-break-after: avoid;
        }}

        .ieee-authors-row {{
            display: grid;
            grid-template-columns: 1fr 1fr 1fr;
            gap: 0.25in;
            margin-bottom: 10px;
            text-align: center;
        }}

        .ieee-author {{
            font-size: 10pt;
            line-height: 1.2;
        }}

        .author-name {{
            font-weight: bold;
            margin-bottom: 3px;
        }}

        .author-affiliation {{
            font-style: italic;
            margin-bottom: 2px;
        }}

        .author-email {{
            font-size: 9pt;
            margin-top: 2px;
        }}

        /* TWO-COLUMN LAYOUT for body content */
        .ieee-two-column {{
            columns: 2;
            column-gap: 0.25in;
            column-fill: balance;
            column-rule: none;
        }}

        /* ABSTRACT and KEYWORDS */
        .ieee-abstract, .ieee-keywords {{
            font-size: 9pt;
            font-weight: bold;
            margin: 15px 0;
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
            break-inside: avoid;
        }}

        /* SECTION HEADINGS - centered, bold, uppercase */
        .ieee-heading {{
            font-size: 10pt;
            font-weight: bold;
            text-align: center;
            text-transform: uppercase;
            margin: 15px 0 5px 0;
            page-break-after: avoid;
            break-after: avoid;
            letter-spacing: 0;
            word-spacing: 0;
        }}

        /* PARAGRAPHS - ultra-aggressive justification */
        .ieee-paragraph {{
            font-size: 10pt;
            margin: 0 0 12px 0;
            text-align: justify !important;
            text-justify: distribute !important;
            text-align-last: justify !important;
            hyphens: auto !important;
            letter-spacing: 0.02em !important;
            word-spacing: 0.12em !important;
            orphans: 2;
            widows: 2;

            /* WeasyPrint specific - ultra-aggressive */
            -weasy-text-align-last: justify !important;
            -weasy-text-justify: distribute !important;
            -weasy-hyphens: auto !important;
        }}

        /* TABLES - exact IEEE formatting */
        .ieee-table-container {{
            margin: 12px 0;
            break-inside: avoid;
            page-break-inside: avoid;
        }}

        .ieee-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 9pt;
            margin: 6px auto;
            border: 1px solid black;
        }}

        .ieee-table-header {{
            border: 1px solid black;
            padding: 4px 6px;
            text-align: center;
            font-weight: bold;
            background-color: #f5f5f5;
            vertical-align: middle;
        }}

        .ieee-table-cell {{
            border: 1px solid black;
            padding: 4px 6px;
            text-align: left;
            vertical-align: top;
        }}

        .ieee-table-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6px 0 12px 0;
            break-before: avoid;
        }}

        /* TABLE NAME - appears before image tables */
        .ieee-table-name {{
            font-size: 9pt;
            font-weight: bold;
            text-align: center;
            margin: 6pt 0 3pt 0;
            text-transform: uppercase;
            letter-spacing: 0.5pt;
        }}

        /* IMAGES - exact sizing and positioning */
        .ieee-image-container {{
            text-align: center;
            margin: 12px 0;
            break-inside: avoid;
            page-break-inside: avoid;
        }}

        .ieee-image {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 0 auto;
        }}

        .ieee-image-very-small {{ width: 1.5in; }}
        .ieee-image-small {{ width: 2.0in; }}
        .ieee-image-medium {{ width: 2.5in; }}
        .ieee-image-large {{ width: 3.3125in; }}

        .ieee-figure-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6px 0 12px 0;
            break-before: avoid;
        }}

        /* REFERENCES */
        .ieee-reference {{
            font-size: 9pt;
            margin: 3px 0;
            padding-left: 15px;
            text-indent: -15px;
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
            letter-spacing: -0.02em;
            word-spacing: 0.05em;
        }}

        /* PAGE BREAKS */
        .page-break {{
            page-break-before: always;
            break-before: page;
        }}

        .keep-together {{
            break-inside: avoid;
            page-break-inside: avoid;
        }}

        /* PRINT OPTIMIZATIONS */
        @media print {{
            body {{
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}

            .ieee-table {{
                border-collapse: collapse !important;
            }}

            .ieee-table-header,
            .ieee-table-cell {{
                border: 1px solid black !important;
            }}
        }}
    </style>
</head>
<body>
    <div class="ieee-title">{title}</div>
    {authors_html}

    <div class="ieee-two-column">
        {f'<div class="ieee-abstract"><strong>Abstract—</strong>{abstract}</div>' if abstract else ''}
        {f'<div class="ieee-keywords"><strong>Index Terms—</strong>{keywords}</div>' if keywords else ''}

        {sections_html}

        {references_html}
    </div>
</body>
</html>"""

    return html


def weasyprint_pdf_from_html(html):
    """Convert HTML to PDF using WeasyPrint with pixel-perfect formatting matching Word"""
    try:
        from weasyprint import CSS, HTML
        from weasyprint.text.fonts import FontConfiguration

        # Create font configuration for better typography
        font_config = FontConfiguration()

        # Minimal CSS overrides - the HTML already contains pixel-perfect CSS
        additional_css = CSS(
            string="""
            /* PDF-specific optimizations */
            @page {
                size: letter;
                margin: 0.75in;
            }

            /* Remove preview-only elements for PDF */
            .preview-note {
                display: none !important;
            }
            
            /* Ensure perfect print rendering */
            body {
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }
            
            /* Force table borders in PDF */
            .ieee-table,
            .ieee-table-header,
            .ieee-table-cell {
                border: 1px solid black !important;
            }
        """
        )

        # Generate PDF with WeasyPrint using the unified HTML
        html_doc = HTML(string=html)
        pdf_bytes = html_doc.write_pdf(
            stylesheets=[additional_css], 
            font_config=font_config, 
            optimize_images=True,
            presentational_hints=True
        )

        print("✅ PDF generated with WeasyPrint - pixel-perfect match to Word", file=sys.stderr)
        return pdf_bytes

    except (ImportError, OSError) as e:
        print(f"⚠️ WeasyPrint not available ({e}), using ReportLab fallback", file=sys.stderr)
        return reportlab_pdf_from_html(html)


def reportlab_pdf_from_html(html):
    """ENHANCED FALLBACK: Convert HTML to PDF using ReportLab with PERFECT justification matching HTML preview"""
    try:
        import re

        from bs4 import BeautifulSoup
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        print(
            "🔧 Using ENHANCED ReportLab with aggressive justification settings...",
            file=sys.stderr,
        )

        # Parse HTML to extract content
        soup = BeautifulSoup(html, "html.parser")

        # Create PDF buffer
        buffer = BytesIO()

        # Create document with IEEE margins
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Create styles for IEEE formatting with ENHANCED justification
        styles = getSampleStyleSheet()

        # IEEE styles with AGGRESSIVE justification settings
        title_style = ParagraphStyle(
            "IEEETitle",
            parent=styles["Title"],
            fontSize=24,
            fontName="Times-Bold",
            alignment=TA_CENTER,
            spaceAfter=20,
        )

        # ENHANCED body style with perfect justification
        body_style = ParagraphStyle(
            "IEEEBody",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Times-Roman",
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leftIndent=0,
            rightIndent=0,
            # AGGRESSIVE justification settings
            wordWrap="LTR",
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=1,
            # Enhanced line spacing for better justification
            leading=12,
            # Character spacing adjustment for perfect justification
            spaceShrinkage=0.10,
            spaceStretchage=0.20,
            # Force full justification with aggressive word spacing
            justifyLastLine=1,
            justifyBreaks=1,
        )

        # ENHANCED abstract style with perfect justification
        abstract_style = ParagraphStyle(
            "IEEEAbstract",
            parent=body_style,
            fontSize=9,
            fontName="Times-Bold",
            alignment=TA_JUSTIFY,
            # Enhanced justification for abstract
            wordWrap="LTR",
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=1,
            leading=10.8,  # 1.2 * fontSize for IEEE standard
        )

        # ENHANCED paragraph style specifically for content blocks
        enhanced_body_style = ParagraphStyle(
            "IEEEBodyEnhanced",
            parent=body_style,
            # CRITICAL: Enhanced justification parameters
            wordWrap="LTR",
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=1,
            # Better line spacing for justification
            leading=12,
            autoLeading="min",
            # Enhanced character spacing
            spaceShrinkage=0.05,
            spaceStretchage=0.15,
        )

        # Build document content
        story = []

        # Extract and add title
        title_elem = soup.find(class_="ieee-title")
        if title_elem:
            story.append(Paragraph(title_elem.get_text(), title_style))
            story.append(Spacer(1, 12))

        # Extract and add authors (simplified)
        authors_elem = soup.find(class_="ieee-authors-container")
        if authors_elem:
            author_names = [
                elem.get_text() for elem in authors_elem.find_all(class_="author-name")
            ]
            if author_names:
                author_text = ", ".join(author_names)
                author_style = ParagraphStyle(
                    "IEEEAuthor",
                    parent=styles["Normal"],
                    fontSize=10,
                    fontName="Times-Roman",
                    alignment=TA_CENTER,
                )
                story.append(Paragraph(author_text, author_style))
                story.append(Spacer(1, 20))

        # Extract and add abstract with ENHANCED justification
        abstract_elem = soup.find(class_="ieee-abstract")
        if abstract_elem:
            # Clean abstract text and apply enhanced formatting
            abstract_text = abstract_elem.get_text()
            story.append(Paragraph(abstract_text, abstract_style))
            story.append(Spacer(1, 12))

        # Extract and add keywords with ENHANCED justification
        keywords_elem = soup.find(class_="ieee-keywords")
        if keywords_elem:
            keywords_text = keywords_elem.get_text()
            story.append(Paragraph(keywords_text, abstract_style))
            story.append(Spacer(1, 20))

        # Extract and add sections with ENHANCED justification
        for heading in soup.find_all(class_="ieee-heading"):
            heading_style = ParagraphStyle(
                "IEEEHeading",
                parent=styles["Heading1"],
                fontSize=10,
                fontName="Times-Bold",
                alignment=TA_CENTER,
                spaceAfter=6,
                spaceBefore=15,
            )
            story.append(Paragraph(heading.get_text(), heading_style))

        # Extract and add paragraphs with PERFECT justification
        for para in soup.find_all(class_="ieee-paragraph"):
            # Use ENHANCED body style for better justification
            para_text = para.get_text()
            # Clean up the text for better ReportLab processing
            para_text = re.sub(r"\s+", " ", para_text).strip()
            story.append(Paragraph(para_text, enhanced_body_style))

        # ENHANCED content block processing
        for content_block in soup.find_all(class_="content-block"):
            block_text = content_block.get_text()
            block_text = re.sub(r"\s+", " ", block_text).strip()
            if block_text:
                story.append(Paragraph(block_text, enhanced_body_style))

        # Extract and add references with ENHANCED justification
        for ref in soup.find_all(class_="ieee-reference"):
            ref_style = ParagraphStyle(
                "IEEEReference",
                parent=enhanced_body_style,
                fontSize=9,
                leftIndent=15,
                firstLineIndent=-15,
                # Enhanced justification for references
                wordWrap="LTR",
                allowWidows=0,
                allowOrphans=0,
                splitLongWords=1,
                leading=10.8,
            )
            ref_text = ref.get_text()
            ref_text = re.sub(r"\s+", " ", ref_text).strip()
            story.append(Paragraph(ref_text, ref_style))

        # Build PDF with enhanced settings
        doc.build(story)

        # Get PDF bytes
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        print(
            "✅ PDF generated with ENHANCED ReportLab - AGGRESSIVE justification achieved",
            file=sys.stderr,
        )
        return pdf_bytes

    except ImportError as e:
        print(f"❌ ReportLab also not available: {e}", file=sys.stderr)
        raise Exception(
            "PDF generation requires WeasyPrint or ReportLab. Both are unavailable."
        )


def pandoc_html_to_docx(html, template_path=None):
    """Convert master HTML to DOCX using pypandoc with IEEE template"""
    try:
        import pypandoc

        # Check if pandoc is available
        try:
            pypandoc.get_pandoc_version()
        except OSError as e:
            print(
                f"⚠️ Pandoc binary not available ({e}), using HTML-to-DOCX converter",
                file=sys.stderr,
            )
            return html_to_docx_converter(html)

        # Create temporary HTML file
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as temp_html:
            temp_html.write(html)
            temp_html_path = temp_html.name

        # Create temporary DOCX file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
            temp_docx_path = temp_docx.name

        try:
            # Convert HTML to DOCX using pypandoc
            extra_args = []
            if template_path and os.path.exists(template_path):
                extra_args.extend(["--reference-doc", template_path])

            # Add additional pandoc options for better formatting
            extra_args.extend(["--standalone", "--wrap=none", "--columns=72"])

            pypandoc.convert_file(
                temp_html_path, "docx", outputfile=temp_docx_path, extra_args=extra_args
            )

            # Read the generated DOCX
            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()

            print(
                "✅ DOCX generated with pypandoc - HTML structure preserved",
                file=sys.stderr,
            )
            return docx_bytes

        finally:
            # Clean up temporary files
            try:
                os.unlink(temp_html_path)
                os.unlink(temp_docx_path)
            except:
                pass

    except ImportError:
        print("⚠️ pypandoc not available, using HTML-to-DOCX converter", file=sys.stderr)
        return html_to_docx_converter(html)
    except Exception as e:
        print(
            f"⚠️ pypandoc conversion failed ({e}), using HTML-to-DOCX converter",
            file=sys.stderr,
        )
        return html_to_docx_converter(html)


def html_to_docx_converter(html):
    """Convert HTML to DOCX using python-docx and BeautifulSoup - no external dependencies"""
    try:
        import re

        from bs4 import BeautifulSoup
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        print(
            "🔄 Converting HTML to DOCX using python-docx converter...", file=sys.stderr
        )

        # Parse HTML
        soup = BeautifulSoup(html, "html.parser")

        # Create new document
        doc = Document()

        # Set document margins (IEEE standard)
        for section in doc.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Process HTML elements
        def process_element(element, parent_doc):
            if element.name == "h1":
                # Title
                para = parent_doc.add_heading(element.get_text().strip(), level=0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if para.runs:
                    para.runs[0].font.name = "Times New Roman"
                    para.runs[0].font.size = Pt(24)
                    para.runs[0].bold = True

            elif element.name == "h2":
                # Section heading
                para = parent_doc.add_heading(element.get_text().strip(), level=1)
                if para.runs:
                    para.runs[0].font.name = "Times New Roman"
                    para.runs[0].font.size = Pt(10)
                    para.runs[0].bold = True

            elif element.name == "p":
                # Paragraph
                text = element.get_text().strip()
                if text:
                    para = parent_doc.add_paragraph(text)
                    if para.runs:
                        para.runs[0].font.name = "Times New Roman"
                        para.runs[0].font.size = Pt(10)

                    # Check for special classes
                    if "ieee-abstract" in element.get("class", []):
                        if para.runs:
                            para.runs[0].font.size = Pt(9)
                            para.runs[0].bold = True
                    elif "ieee-keywords" in element.get("class", []):
                        if para.runs:
                            para.runs[0].font.size = Pt(9)
                            para.runs[0].bold = True
                    elif "ieee-reference" in element.get("class", []):
                        if para.runs:
                            para.runs[0].font.size = Pt(9)

            elif element.name == "div":
                # Process div contents
                for child in element.children:
                    if hasattr(child, "name"):
                        process_element(child, parent_doc)

        # Process body content
        body = soup.find("body")
        if body:
            for child in body.children:
                if hasattr(child, "name"):
                    process_element(child, doc)

        # Generate DOCX bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()

        print(
            f"✅ HTML-to-DOCX conversion completed: {len(docx_bytes)} bytes",
            file=sys.stderr,
        )
        return docx_bytes

    except Exception as e:
        print(f"❌ HTML-to-DOCX converter failed: {e}", file=sys.stderr)
        return None


def generate_ieee_pdf_perfect_justification(form_data):
    """Generate IEEE-formatted PDF with PERFECT text justification using WeasyPrint - bypasses Word's weak justification"""

    # Extract document data
    title = sanitize_text(form_data.get("title", "Untitled Document"))
    authors = form_data.get("authors", [])
    abstract = sanitize_text(form_data.get("abstract", ""))
    keywords = sanitize_text(form_data.get("keywords", ""))
    sections = form_data.get("sections", [])
    references = form_data.get("references", [])

    # Format authors in IEEE style
    authors_html = ""
    if authors:
        num_authors = len(authors)
        authors_html = '<div class="ieee-authors-container">'

        # Process authors in groups of 3 (IEEE standard)
        authors_per_row = 3
        for row_start in range(0, num_authors, authors_per_row):
            row_end = min(row_start + authors_per_row, num_authors)
            row_authors = authors[row_start:row_end]

            authors_html += '<div class="ieee-authors-row">'
            for author in row_authors:
                author_name = sanitize_text(author.get("name", ""))
                author_html = f'<div class="ieee-author"><strong>{author_name}</strong>'

                # Add affiliation fields
                fields = ["department", "organization", "city", "state", "country"]
                for field in fields:
                    if author.get(field):
                        author_html += f"<br><em>{sanitize_text(author[field])}</em>"

                # Add email
                if author.get("email"):
                    author_html += f'<br><em>{sanitize_text(author["email"])}</em>'

                # Fallback to affiliation field
                if not any(author.get(field) for field in fields) and author.get(
                    "affiliation"
                ):
                    author_html += (
                        f'<br><em>{sanitize_text(author["affiliation"])}</em>'
                    )

                author_html += "</div>"
                authors_html += author_html

            authors_html += "</div>"
        authors_html += "</div>"

    # Process sections with content blocks (tables and images)
    sections_html = ""
    for i, section in enumerate(sections, 1):
        section_title = sanitize_text(section.get("title", ""))
        if section_title:
            sections_html += (
                f'<div class="ieee-heading">{i}. {section_title.upper()}</div>'
            )

        # Process content blocks
        content_blocks = section.get("contentBlocks", [])
        table_count = 0
        img_count = 0

        for block in content_blocks:
            block_type = block.get("type", "text")

            if block_type == "text" and block.get("content"):
                content = sanitize_text(block["content"])
                sections_html += f'<div class="ieee-paragraph">{content}</div>'

            elif block_type == "table":
                table_count += 1
                table_type = block.get("tableType", "interactive")

                if table_type == "interactive":
                    headers = block.get("headers", [])
                    rows_data = block.get("tableData", [])

                    if headers and rows_data:
                        sections_html += '<div class="ieee-table-container">'
                        sections_html += '<table class="ieee-table">'

                        # Header row
                        sections_html += "<thead><tr>"
                        for header in headers:
                            sections_html += f"<th>{sanitize_text(str(header))}</th>"
                        sections_html += "</tr></thead>"

                        # Data rows
                        sections_html += "<tbody>"
                        for row_data in rows_data:
                            sections_html += "<tr>"
                            for cell_data in row_data:
                                sections_html += (
                                    f"<td>{sanitize_text(str(cell_data))}</td>"
                                )
                            sections_html += "</tr>"
                        sections_html += "</tbody>"

                        sections_html += "</table>"

                        # Table caption
                        caption = block.get("caption", block.get("tableName", ""))
                        if caption:
                            sections_html += f'<div class="ieee-table-caption">TABLE {i}.{table_count}: {sanitize_text(caption).upper()}</div>'

                        sections_html += "</div>"

                elif table_type == "image" and block.get("data"):
                    # Handle image tables
                    image_data = block["data"]
                    if "," in image_data:
                        image_data = image_data.split(",")[1]

                    sections_html += f'<div class="ieee-image-container">'
                    sections_html += f'<img src="data:image/png;base64,{image_data}" class="ieee-image ieee-image-{block.get("size", "medium")}" />'

                    # Fix table caption duplication
                    caption_text = block.get("caption", "").strip()
                    table_name = block.get("tableName", "").strip()

                    if caption_text and table_name:
                        if table_name.lower() in caption_text.lower():
                            final_caption = caption_text
                        elif caption_text.lower() in table_name.lower():
                            final_caption = table_name
                        else:
                            final_caption = caption_text
                    else:
                        final_caption = caption_text or table_name

                    if final_caption:
                        sections_html += f'<div class="ieee-table-caption">TABLE {i}.{table_count}: {sanitize_text(final_caption).upper()}</div>'

                    sections_html += "</div>"

            elif block_type == "image" and block.get("data") and block.get("caption"):
                img_count += 1
                image_data = block["data"]
                if "," in image_data:
                    image_data = image_data.split(",")[1]

                sections_html += f'<div class="ieee-image-container">'
                sections_html += f'<img src="data:image/png;base64,{image_data}" class="ieee-image ieee-image-{block.get("size", "medium")}" />'
                sections_html += f'<div class="ieee-figure-caption">FIG. {i}.{img_count}: {sanitize_text(block["caption"]).upper()}</div>'
                sections_html += "</div>"

    # Process references
    references_html = ""
    if references:
        references_html = '<div class="ieee-heading">REFERENCES</div>'
        for i, ref in enumerate(references, 1):
            ref_text = (
                sanitize_text(ref.get("text", ""))
                if isinstance(ref, dict)
                else sanitize_text(str(ref))
            )
            if ref_text:
                references_html += f'<div class="ieee-reference">[{i}] {ref_text}</div>'

    # Create enhanced HTML with PERFECT justification CSS
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        @page {{
            margin: 0.75in;
            size: letter;
            @bottom-center {{
                content: counter(page);
                font-family: 'Times New Roman', serif;
                font-size: 10pt;
            }}
        }}

        body {{
            font-family: 'Times New Roman', serif;
            font-size: 10pt;
            line-height: 1.2;
            margin: 0;
            padding: 0;
            background: white;
            color: black;

            /* PERFECT JUSTIFICATION - LaTeX quality */
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
            -webkit-hyphens: auto;
            -moz-hyphens: auto;
            -ms-hyphens: auto;

            /* Fine-tune character and word spacing for perfect justification */
            letter-spacing: -0.02em;
            word-spacing: 0.05em;

            /* Prevent orphans and widows */
            orphans: 2;
            widows: 2;
        }}

        /* Title formatting */
        .ieee-title {{
            font-size: 24pt;
            font-weight: bold;
            text-align: center;
            margin: 0 0 20px 0;
            line-height: 1.3;
            page-break-after: avoid;
        }}

        /* Authors formatting */
        .ieee-authors-container {{
            text-align: center;
            margin: 15px 0 20px 0;
            page-break-after: avoid;
        }}

        .ieee-authors-row {{
            display: flex;
            justify-content: center;
            margin-bottom: 10px;
        }}

        .ieee-author {{
            flex: 1;
            max-width: 33.33%;
            padding: 0 10px;
            font-size: 10pt;
        }}

        /* Two-column layout for body content */
        .ieee-two-column {{
            columns: 2;
            column-gap: 0.25in;
            column-fill: balance;
        }}

        /* Abstract and keywords */
        .ieee-abstract, .ieee-keywords {{
            margin: 15px 0;
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
            font-size: 9pt;
            font-weight: bold;
        }}

        /* Section headings */
        .ieee-heading {{
            font-weight: bold;
            margin: 15px 0 5px 0;
            text-transform: uppercase;
            font-size: 10pt;
            text-align: center;
            page-break-after: avoid;
        }}

        /* Paragraphs */
        .ieee-paragraph {{
            margin: 0 0 12px 0;
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
            letter-spacing: -0.02em;
            word-spacing: 0.05em;
            orphans: 2;
            widows: 2;
        }}

        /* Tables */
        .ieee-table-container {{
            margin: 12px 0;
            page-break-inside: avoid;
        }}

        .ieee-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 9pt;
            margin: 6px auto;
        }}

        .ieee-table th, .ieee-table td {{
            border: 1px solid black;
            padding: 4px 6px;
            text-align: left;
        }}

        .ieee-table th {{
            font-weight: bold;
            text-align: center;
            background-color: #f5f5f5;
        }}

        .ieee-table-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6px 0 12px 0;
        }}

        /* Images */
        .ieee-image-container {{
            text-align: center;
            margin: 12px 0;
            page-break-inside: avoid;
        }}

        .ieee-image {{
            max-width: 100%;
            height: auto;
        }}

        .ieee-image-small {{ width: 2.0in; }}
        .ieee-image-medium {{ width: 2.5in; }}
        .ieee-image-large {{ width: 3.3125in; }}
        .ieee-image-very-small {{ width: 1.5in; }}

        .ieee-figure-caption {{
            text-align: center;
            font-size: 9pt;
            font-weight: bold;
            margin: 6px 0 12px 0;
        }}

        /* References */
        .ieee-reference {{
            margin: 3px 0;
            padding-left: 15px;
            text-indent: -15px;
            font-size: 9pt;
            text-align: justify;
            text-justify: inter-word;
            hyphens: auto;
        }}

        /* Page breaks */
        .page-break {{
            page-break-before: always;
        }}

        /* Prevent breaking between elements */
        .keep-together {{
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
    <div class="ieee-title">{title}</div>
    {authors_html}

    <div class="ieee-two-column">
        {f'<div class="ieee-abstract"><strong>Abstract—</strong>{abstract}</div>' if abstract else ''}
        {f'<div class="ieee-keywords"><strong>Index Terms—</strong>{keywords}</div>' if keywords else ''}

        {sections_html}

        {references_html}
    </div>
</body>
</html>"""

    try:
        # Try to import and use WeasyPrint for perfect PDF generation
        from weasyprint import CSS, HTML
        from weasyprint.text.fonts import FontConfiguration

        # Create font configuration for better typography
        font_config = FontConfiguration()

        # Additional CSS for even better justification
        additional_css = CSS(
            string="""
            @page {
                margin: 0.75in;
                size: letter;
            }

            body {
                text-rendering: optimizeLegibility;
                font-variant-ligatures: common-ligatures;
                font-feature-settings: "liga" 1, "kern" 1;
            }

            /* FORCE FULL JUSTIFICATION - Match Word document justification */
            .ieee-paragraph, .ieee-abstract, .ieee-keywords, .ieee-reference, p {
                text-align: justify !important;
                text-align-last: justify !important;
                text-justify: inter-word !important;
                hyphens: auto !important;
                word-break: normal;
                overflow-wrap: break-word;
                line-height: 1.2 !important;
            }
        """
        )

        # Generate PDF with WeasyPrint
        html_doc = HTML(string=html)
        pdf_bytes = html_doc.write_pdf(
            stylesheets=[additional_css], font_config=font_config, optimize_images=True
        )

        print(
            "✅ PDF generated with WeasyPrint - perfect justification achieved",
            file=sys.stderr,
        )
        return pdf_bytes

    except (ImportError, OSError) as e:
        print(
            f"⚠️ WeasyPrint not available ({e}), using ReportLab for PDF generation",
            file=sys.stderr,
        )

        # Fallback: Use ReportLab for better PDF generation with justification
        try:
            import io

            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.pdfgen import canvas
            from reportlab.platypus import (
                Image,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )

            # Create PDF buffer
            buffer = BytesIO()

            # Create document with IEEE margins
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            # Create styles for IEEE formatting
            styles = getSampleStyleSheet()

            # IEEE Title style
            title_style = ParagraphStyle(
                "IEEETitle",
                parent=styles["Title"],
                fontSize=24,
                fontName="Times-Bold",
                alignment=TA_CENTER,
                spaceAfter=20,
            )

            # IEEE Body style with perfect justification
            body_style = ParagraphStyle(
                "IEEEBody",
                parent=styles["Normal"],
                fontSize=10,
                fontName="Times-Roman",
                alignment=TA_JUSTIFY,
                spaceAfter=12,
                leftIndent=0,
                rightIndent=0,
                wordWrap="LTR",
            )

            # IEEE Abstract style
            abstract_style = ParagraphStyle(
                "IEEEAbstract",
                parent=body_style,
                fontSize=9,
                fontName="Times-Bold",
                alignment=TA_JUSTIFY,
            )

            # Build document content
            story = []

            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

            # Add authors (simplified for ReportLab)
            if authors:
                author_text = ", ".join([author.get("name", "") for author in authors])
                author_style = ParagraphStyle(
                    "IEEEAuthor",
                    parent=styles["Normal"],
                    fontSize=10,
                    fontName="Times-Roman",
                    alignment=TA_CENTER,
                )
                story.append(Paragraph(author_text, author_style))
                story.append(Spacer(1, 20))

            # Add abstract
            if abstract:
                story.append(Paragraph(f"<b>Abstract—</b>{abstract}", abstract_style))
                story.append(Spacer(1, 12))

            # Add keywords
            if keywords:
                story.append(
                    Paragraph(f"<b>Index Terms—</b>{keywords}", abstract_style)
                )
                story.append(Spacer(1, 20))

            # Add sections
            for i, section in enumerate(sections, 1):
                section_title = sanitize_text(section.get("title", ""))
                if section_title:
                    heading_style = ParagraphStyle(
                        "IEEEHeading",
                        parent=styles["Heading1"],
                        fontSize=10,
                        fontName="Times-Bold",
                        alignment=TA_CENTER,
                        spaceAfter=6,
                        spaceBefore=15,
                    )
                    story.append(
                        Paragraph(f"{i}. {section_title.upper()}", heading_style)
                    )

                # Process content blocks
                content_blocks = section.get("contentBlocks", [])
                for block in content_blocks:
                    if block.get("type") == "text" and block.get("content"):
                        content = sanitize_text(block["content"])
                        story.append(Paragraph(content, body_style))

            # Add references
            if references:
                heading_style = ParagraphStyle(
                    "IEEEHeading",
                    parent=styles["Heading1"],
                    fontSize=10,
                    fontName="Times-Bold",
                    alignment=TA_CENTER,
                    spaceAfter=6,
                    spaceBefore=15,
                )
                story.append(Paragraph("REFERENCES", heading_style))

                ref_style = ParagraphStyle(
                    "IEEEReference",
                    parent=body_style,
                    fontSize=9,
                    leftIndent=15,
                    firstLineIndent=-15,
                )

                for i, ref in enumerate(references, 1):
                    ref_text = (
                        sanitize_text(ref.get("text", ""))
                        if isinstance(ref, dict)
                        else sanitize_text(str(ref))
                    )
                    if ref_text:
                        story.append(Paragraph(f"[{i}] {ref_text}", ref_style))

            # Build PDF
            doc.build(story)

            # Get PDF bytes
            buffer.seek(0)
            pdf_bytes = buffer.getvalue()
            buffer.close()

            print(
                "✅ PDF generated with ReportLab - good justification achieved",
                file=sys.stderr,
            )
            return pdf_bytes

        except ImportError as reportlab_error:
            print(
                f"❌ ReportLab also not available: {reportlab_error}", file=sys.stderr
            )
            raise Exception(
                "PDF generation requires WeasyPrint or ReportLab. Both are unavailable."
            )


def main():
    """Main function with unified rendering system for pixel-perfect DOCX/PDF matching."""

    # Parse command line arguments
    parser = argparse.ArgumentParser(
        description="IEEE Document Generator with unified rendering system"
    )
    parser.add_argument(
        "--debug-compare",
        action="store_true",
        help="Generate both DOCX and PDF for visual comparison",
    )
    parser.add_argument(
        "--output",
        choices=["docx", "pdf", "html"],
        default="docx",
        help="Output format (default: docx)",
    )

    # Parse args if running from command line, otherwise use defaults
    if len(sys.argv) > 1:
        args = parser.parse_args()
    else:
        args = argparse.Namespace(debug_compare=False, output="docx")

    try:
        # Read JSON data from stdin
        input_data = sys.stdin.read()
        form_data = json.loads(input_data)

        # Override output type from form data if present
        output_type = form_data.get("output", args.output).lower()

        # Build unified document model with exact OpenXML formatting metadata
        print("🎯 Building unified document model with pixel-perfect formatting...", file=sys.stderr)
        model = build_document_model(form_data)
        print("✅ Document model built - single source of truth for all formats", file=sys.stderr)

        if args.debug_compare:
            # DEBUG MODE: Generate both formats for comparison
            print("🔍 DEBUG MODE: Generating both DOCX and PDF for visual comparison...", file=sys.stderr)

            # Generate DOCX using original perfect generator (unchanged)
            print("📄 Generating DOCX using original perfect generator...", file=sys.stderr)
            docx_bytes = generate_ieee_document(form_data)

            # Generate HTML using unified rendering
            print("🌐 Generating HTML using unified rendering system...", file=sys.stderr)
            html = render_to_html(model)

            # Generate PDF from unified HTML
            print("🎯 Generating PDF from unified HTML...", file=sys.stderr)
            pdf_bytes = weasyprint_pdf_from_html(html)

            # Save all files for comparison
            import time
            timestamp = str(int(time.time()))

            with open(f"debug_compare_{timestamp}.docx", "wb") as f:
                f.write(docx_bytes)
            print(f"📁 DOCX saved: debug_compare_{timestamp}.docx", file=sys.stderr)

            with open(f"debug_compare_{timestamp}.html", "w", encoding="utf-8") as f:
                f.write(html)
            print(f"📁 HTML saved: debug_compare_{timestamp}.html", file=sys.stderr)

            with open(f"debug_compare_{timestamp}.pdf", "wb") as f:
                f.write(pdf_bytes)
            print(f"📁 PDF saved: debug_compare_{timestamp}.pdf", file=sys.stderr)

            print("🔍 Open all files to verify pixel-perfect matching", file=sys.stderr)

            # Return the requested format
            if output_type == "pdf":
                doc_data = pdf_bytes
            elif output_type == "html":
                doc_data = html.encode('utf-8')
            else:
                doc_data = docx_bytes

        elif output_type == "pdf":
            # Generate PDF using unified rendering system
            print("🎯 Generating PDF using unified rendering system...", file=sys.stderr)
            html = render_to_html(model)
            doc_data = weasyprint_pdf_from_html(html)
            print("✅ PDF generated with pixel-perfect formatting matching Word", file=sys.stderr)

        elif output_type == "html":
            # Generate HTML preview using unified rendering system
            print("🌐 Generating HTML preview using unified rendering system...", file=sys.stderr)
            html = render_to_html(model)
            
            # Add preview note for live preview
            preview_note = '''
    <div style="background: #e8f4fd; border: 1px solid #bee5eb; padding: 12px; margin: 20px 0; font-size: 9pt; color: #0c5460; text-align: center; border-radius: 4px;">
        📄 IEEE Live Preview - This is exactly what your PDF will look like
    </div>
    '''
            html = html.replace('<body>', f'<body>{preview_note}', 1)
            
            doc_data = html.encode('utf-8')
            print("✅ HTML preview generated with pixel-perfect formatting", file=sys.stderr)

        else:
            # Generate DOCX using original perfect generator (unchanged)
            print("📄 Generating DOCX using original perfect generator...", file=sys.stderr)
            doc_data = generate_ieee_document(form_data)
            print("✅ DOCX generated with perfect IEEE formatting", file=sys.stderr)

        # Write data to stdout
        if output_type == "html":
            sys.stdout.write(doc_data.decode('utf-8'))
        else:
            sys.stdout.buffer.write(doc_data)

    except Exception as e:
        import traceback

        sys.stderr.write(f"Error: {str(e)}\n")
        sys.stderr.write(f"Traceback: {traceback.format_exc()}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
