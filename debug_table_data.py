#!/usr/bin/env python3
"""
Debug script to specifically test table data processing in DOCX
"""

import json
import sys
import os

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from ieee_generator_fixed import build_document_model, render_to_html, generate_ieee_document

def test_table_data():
    """Test table data processing specifically"""
    print("🧪 TESTING TABLE DATA PROCESSING")
    print("=" * 50)
    
    # Create test data with ONLY a table
    test_data = {
        "title": "Table Test Document",
        "authors": [
            {"name": "Test Author", "department": "Computer Science", "email": "test@university.edu"}
        ],
        "abstract": "This document tests table processing.",
        "keywords": "tables, IEEE format",
        "sections": [
            {
                "id": "section1",
                "title": "Table Test Section",
                "contentBlocks": [
                    {
                        "id": "table1",
                        "type": "table",
                        "tableType": "interactive",
                        "tableName": "Performance Data",
                        "caption": "System Performance Results",
                        "headers": ["Algorithm", "Speed (ms)", "Accuracy (%)"],
                        "tableData": [
                            ["Method A", "100", "95.5"],
                            ["Method B", "150", "97.2"],
                            ["Method C", "80", "93.8"]
                        ],
                        "size": "medium",
                        "order": 1
                    }
                ],
                "order": 1
            }
        ],
        "references": []
    }
    
    try:
        print("📋 Building document model...")
        model = build_document_model(test_data)
        
        print("📊 Analyzing model structure:")
        sections = model.get("sections", [])
        
        for i, section in enumerate(sections):
            content_blocks = section.get("content_blocks", [])
            print(f"  Section {i+1}: '{section.get('title', 'Untitled')}'")
            print(f"    - Content blocks: {len(content_blocks)}")
            
            for j, block in enumerate(content_blocks):
                block_type = block.get("type", "unknown")
                print(f"      Block {j+1}: {block_type}")
                
                if block_type == "table":
                    print(f"        - Table type: {block.get('table_type', 'not set')}")
                    print(f"        - Headers: {block.get('headers', 'not set')}")
                    print(f"        - Table data: {block.get('table_data', 'not set')}")
                    print(f"        - Caption: {block.get('caption', {}).get('text', 'not set')}")
                    print(f"        - Raw block data:")
                    for key, value in block.items():
                        if key not in ['table_data']:
                            print(f"          {key}: {value}")
                        else:
                            print(f"          {key}: {len(value) if value else 0} rows")
        
        print("\n🌐 Rendering to HTML...")
        html = render_to_html(model)
        
        # Count table elements in HTML
        table_count = html.count('<table')
        table_rows = html.count('<tr>')
        table_cells = html.count('<td>')
        
        print(f"📊 HTML Analysis:")
        print(f"  • HTML tables: {table_count}")
        print(f"  • HTML rows: {table_rows}")
        print(f"  • HTML cells: {table_cells}")
        
        # Save HTML for inspection
        html_file = "debug_table_data_output.html"
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"📁 HTML saved: {html_file}")
        
        print("\n📄 Generating DOCX...")
        docx_bytes = generate_ieee_document(test_data)
        
        docx_file = "debug_table_data_output.docx"
        with open(docx_file, 'wb') as f:
            f.write(docx_bytes)
        print(f"📁 DOCX saved: {docx_file} ({len(docx_bytes)} bytes)")
        
        # Try to analyze the DOCX content
        try:
            from docx import Document
            doc = Document(docx_file)
            
            print(f"\n📄 DOCX Analysis:")
            print(f"  • Total paragraphs: {len(doc.paragraphs)}")
            print(f"  • Total tables: {len(doc.tables)}")
            
            for i, table in enumerate(doc.tables):
                print(f"    Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(f"      Row {row_idx+1}: {row_data}")
            
            print(f"  • Paragraph contents:")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    print(f"    Para {i+1}: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        except Exception as e:
            print(f"❌ Error analyzing DOCX: {e}")
        
        print(f"\n🎉 TEST COMPLETE!")
        print("=" * 50)
        
        return True
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Starting table data test...")
    success = test_table_data()
    if success:
        print("✅ Test completed successfully")
    else:
        print("❌ Test failed")
        sys.exit(1)